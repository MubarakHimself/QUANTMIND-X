"""
Generate QuantMindX Planning Addendum — Session 2 (March 2026)
Addendum to: QuantMindX_Planning_Document_March2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex=None):
    p = doc.add_heading(text, level=level)
    if color_hex:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def add_body(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    return p

def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("NOTE: " + text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def make_table(doc, headers, rows, header_bg="1F3864", col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            cell.text = val
            set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

# ─── Build Document ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("QUANTMINDX — ITT")
tr.font.size = Pt(22)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Planning Addendum — Session 2 | March 2026")
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
ref_p = doc.add_paragraph()
ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ref_r = ref_p.add_run(
    "Addendum to: QuantMindX_Planning_Document_March2026.docx\n"
    "Cross-references: _bmad-output/planning-artifacts/prd.md | architecture.md | epics.md"
)
ref_r.font.size = Pt(9.5)
ref_r.font.italic = True
ref_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

p_scope = doc.add_paragraph()
p_scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
ps_r = p_scope.add_run(
    "This document contains ONLY new architectural decisions, frameworks, and design additions\n"
    "arising from Session 2. It does not repeat existing planning documents."
)
ps_r.font.size = Pt(9.5)
ps_r.font.bold = True
ps_r.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

doc.add_page_break()

# ── SECTION 1: STRATEGIC PRINCIPLES ─────────────────────────────────────────
add_heading(doc, "1. Strategic Principles (Confirmed This Session)", level=1, color_hex="1F3864")

add_body(doc,
    "The following principles govern all new architectural decisions in this addendum. "
    "They are not repeated from the prior planning document but are reaffirmed here as "
    "the foundation for every new design choice.")

add_heading(doc, "1.1 The Edge Is Probabilistic Scale, Not Speed", level=2, color_hex="2E5496")
add_body(doc,
    "QUANTMINDX does not compete on co-location or hardware latency. The system's edge is: "
    "high win-rate × correct Kelly sizing × many concurrent variants × fast agentic iteration = "
    "sustainable compounding from a small book. Speed is a secondary concern after signal quality.")

add_heading(doc, "1.2 No AI in the Live Execution Hot Path", level=2, color_hex="2E5496")
add_body(doc,
    "LLM inference has no place in the Cloudzy execution path. Alpha Forge, research, and "
    "improvement loops are Contabo workloads. Cloudzy runs: Sentinel → Governor → Commander → ZMQ. "
    "All cached, pre-computed, async. No LLM call between signal and order.")

add_heading(doc, "1.3 Trust the Validated EA", level=2, color_hex="2E5496")
add_body(doc,
    "Any bot that has passed AlphaForge Workflow 1, full Workflow 2 (all six backtest modes), "
    "paper trading gate, and human approval carries a validated edge. The system does NOT override "
    "that EA's trade logic on normal regime cycling. Layer 3 interventions are reserved exclusively "
    "for CHAOS conditions, prop firm drawdown breaches, and kill switch activation.")

add_heading(doc, "1.4 Session Mask Is the Runtime Constraint", level=2, color_hex="2E5496")
add_body(doc,
    "The system runs as long as liquidity windows are active. Intraday P&L lock mode is explicitly "
    "removed. The natural stops are: session mask end, 3/5/7 drawdown limits (existing risk framework), "
    "and circuit breaker thresholds. Nothing else pauses the system.")

doc.add_paragraph()

# ── SECTION 2: BROKER ARCHITECTURE ──────────────────────────────────────────
add_heading(doc, "2. Broker Architecture", level=1, color_hex="1F3864")

add_heading(doc, "2.1 Current State — Exness Raw Spread", level=2, color_hex="2E5496")
add_body(doc,
    "Current broker: Exness Raw Spread account (Pro tier). Raw spread (0.0 pips + fixed commission) "
    "provides a clean, calculable gross edge for the Kelly formula — no hidden spread drag corrupting "
    "the edge estimate. This is the direct mechanism behind the House of Money effect functioning "
    "correctly: when the Kelly engine operates on a known-cost ECN feed, house money scaling kicks "
    "in on an accurate Kelly base rather than a degraded one.")
add_body(doc,
    "Exness ECN volume through MT5 is real lot-based volume from their liquidity pool — a valid proxy "
    "for true market volume on major pairs. Adequate for Phase 1.")

add_heading(doc, "2.2 Target State — IC Markets ECN", level=2, color_hex="2E5496")
add_body(doc,
    "IC Markets (LD4 London) is the target broker upgrade. Reasons: deeper ECN liquidity pool, "
    "confirmed real contract volume (not just tick count) for VWAP/Volume Profile accuracy, "
    "explicit scalping + algorithmic trading approval, MT5 support, Islamic account availability, "
    "and LD4 server proximity matching Cloudzy's deployment region for 1–3ms tick arrival latency.")
add_note(doc,
    "Migration prerequisite: Eversend bank account setup for Uganda-based onboarding. "
    "Migration does not require code changes — only broker credentials in .env.")

add_heading(doc, "2.3 MT5 Confirmed — Magic Number Architecture", level=2, color_hex="2E5496")
add_body(doc,
    "MetaTrader 5 is confirmed as the execution platform. No platform migration unless a competing "
    "platform offers demonstrably better data AND execution. The multi-bot architecture uses MT5's "
    "native Magic Number system: every EA carries a unique integer identifier stamped on every order "
    "it places. One broker account, one balance — 200+ magic numbers providing full trade attribution.")
make_table(doc,
    ["Magic Number Range", "Pool", "Account", "Strategy Family"],
    [
        ["1000–4999", "Long Scalper Pool", "Machine Gun (RoboForex/IC Markets)", "Scalping — long bias"],
        ["5000–8999", "Short Scalper Pool", "Machine Gun (RoboForex/IC Markets)", "Scalping — short bias"],
        ["9000–11999", "Neutral/Range Pool", "Machine Gun", "Scalping — mean reversion"],
        ["12000–15999", "ORB Long Pool", "Sniper (Exness Raw → IC Markets)", "ORB — directional long"],
        ["16000–19999", "ORB Short Pool", "Sniper (Exness Raw → IC Markets)", "ORB — directional short"],
        ["20000–20999", "ORB False Breakout Pool", "Sniper", "ORB — false breakout reversal"],
    ],
    col_widths=[1.5, 1.6, 2.2, 2.0]
)

doc.add_paragraph()

# ── SECTION 3: EXECUTION LAYER — SL/TP THREE-LAYER ARCHITECTURE ─────────────
add_heading(doc, "3. Execution Layer: Three-Layer SL/TP Architecture (New)", level=1, color_hex="1F3864")

add_body(doc,
    "This architecture is new and not specified in the prior planning document. It defines "
    "how stop loss and take profit are managed across three independent layers.")

add_heading(doc, "Layer 1 — Hard Safety SL/TP (EA at Broker Level)", level=2, color_hex="2E5496")
add_body(doc,
    "Every EA sets a hard SL and TP at order placement time via MQL5. This is non-negotiable "
    "and must be present in every EA template. This is the safety net: if Python crashes, "
    "ZMQ disconnects, or Cloudzy goes offline, the broker still holds the stop. "
    "For scalping: hard SL = 1.5–2× ATR for that timeframe. Hard TP = 1.2–1.5× the stop distance.")
add_note(doc, "This is already implied by the EA tag architecture in epics.md but must be "
         "explicitly enforced in every generated EA template during AlphaForge Workflow 1.")

add_heading(doc, "Layer 2 — Dynamic Position Modification (Cloudzy-Local Only)", level=2, color_hex="2E5496")
add_body(doc,
    "A lightweight Position Monitor service runs on Cloudzy. It holds a table of open positions "
    "and monitors two conditions per position: (1) price has moved 1R in favour → move SL to "
    "break-even, and (2) Sentinel cached regime has changed (published to Redis). When either "
    "condition is met, the Position Monitor sends TRADE_ACTION_MODIFY directly via ZMQ.")
add_body(doc,
    "Layer 2 MUST run on Cloudzy — not Contabo. Routing a stop modification through Contabo → "
    "Cloudzy → ZMQ adds 5–15ms of inter-node HTTP latency before execution begins. This is "
    "unacceptable for an open position on a scalping timeframe.")

add_labeled(doc, "Layer 2 latency budget", "Redis read (1ms) + ZMQ modify (5–20ms broker RT) = 6–21ms total")

add_body(doc, "Multi-timeframe conflict resolution for Layer 2 decisions:", bold=True)
add_bullet(doc, "Only M5 regime shifts, H1 intact → no action. M5 noise is expected. H1 still valid.")
add_bullet(doc, "H1 flips against open position → move SL to nearest structural level (tighten, not close).")
add_bullet(doc, "Both H1 AND H4 flip against open position → tighten aggressively. Accept reduced loss.")
add_bullet(doc, "CHAOS signal from Lyapunov or Correlation Sensor → Layer 3 applies (see Section 3 Layer 3).")

add_heading(doc, "Regime Persistence Timer (New Component)", level=3, color_hex="375623")
add_body(doc,
    "Before any Layer 2 action is triggered by a regime shift, the new regime must persist "
    "for N consecutive bars (configurable per timeframe). Default: 3 M5 bars = 15 minutes for "
    "scalping, 2 H1 bars = 2 hours for ORB. If the regime reverts before N bars, the timer "
    "resets and no Layer 2 action is taken. This prevents false regime flips from whipsaw "
    "markets triggering unnecessary stop modifications.")

add_heading(doc, "Layer 3 — System-Level Forced Exit (CHAOS + Kill Switch Only)", level=2, color_hex="2E5496")
add_body(doc,
    "The validated EA's trade logic is trusted for normal market conditions. Layer 3 applies "
    "exclusively in three scenarios:")
add_bullet(doc, "CHAOS condition: Lyapunov exponent above threshold OR Correlation Sensor max_eigenvalue > 2.0. "
           "These represent conditions outside the backtest envelope — not normal regime cycling.")
add_bullet(doc, "Prop firm drawdown breach: PropFirmOverlay hard limit reached.")
add_bullet(doc, "Kill switch activation: Any tier (soft/progressive/full) triggered by operator.")
add_body(doc, "For ORB open positions under Layer 3: tighten stop to lock partial profit rather than "
         "immediately closing. Immediate close only on hard CHAOS or kill switch. For scalping: "
         "immediate close is acceptable given the short holding period and small per-trade R.")

doc.add_paragraph()

# ── SECTION 4: REGIME-CONDITIONAL STRATEGY POOL FRAMEWORK ───────────────────
add_heading(doc, "4. Regime-Conditional Strategy Pool Framework (New)", level=1, color_hex="1F3864")

add_body(doc,
    "Replaces the concept of a single bot with a 'directional flip switch.' The system maintains "
    "three competing pools for scalping and two for ORB, all activated by Sentinel regime state.")

add_heading(doc, "4.1 Scalping Pools", level=2, color_hex="2E5496")
make_table(doc,
    ["Pool", "Activation Condition", "Regime Confidence Threshold", "Direction"],
    [
        ["Long Scalper Pool", "H1/H4 Sentinel = TREND_UP", "> 0.65", "Long entries only"],
        ["Short Scalper Pool", "H1/H4 Sentinel = TREND_DOWN", "> 0.65", "Short entries only"],
        ["Neutral/Range Pool", "RANGE_STABLE or confidence < 0.65", "N/A", "Mean-reversion both directions"],
    ],
    col_widths=[1.5, 2.4, 1.8, 1.6]
)
add_body(doc, "When regime confidence is 0.40–0.65: only Neutral/Range pool activates. "
         "When regime confidence < 0.40 or regime = HIGH_CHAOS: no new scalp entries. "
         "All open positions managed by their own SL/TP (trust the EA).")

add_heading(doc, "4.2 ORB Pools", level=2, color_hex="2E5496")
make_table(doc,
    ["Pool", "Activation Condition", "Notes"],
    [
        ["ORB Long Pool", "Session open breakout above range high + volume confirmed + H1 TREND_UP or BREAKOUT_PRIME", "Primary ORB direction"],
        ["ORB Short Pool", "Session open breakout below range low + volume confirmed", "Primary ORB direction"],
        ["ORB False Breakout Pool", "Breakout fires but price reverses back through range within 2 bars + volume surge", "Reversal variant — own signal, own backtest history"],
    ],
    col_widths=[1.8, 3.2, 1.7]
)
add_note(doc, "The False Breakout Pool is a distinct strategy family with its own backtest history. "
         "It is NOT the ORB bot with a direction toggle — it is trained specifically on the false "
         "breakout pattern and is only active when that specific setup criteria is met.")

add_heading(doc, "4.3 Routing Matrix Adjustment", level=2, color_hex="2E5496")
add_body(doc,
    "The routing matrix (src/router/routing_matrix.py) must be updated to reflect the two "
    "primary strategy families explicitly: Scalping family (Machine Gun account) and ORB family "
    "(Sniper account). StrategyType enum to add: SCALPER_LONG, SCALPER_SHORT, SCALPER_RANGE, "
    "ORB_LONG, ORB_SHORT, ORB_FALSE_BREAKOUT. Commander routing logic gates pool activation "
    "against Sentinel regime confidence before dispatching.")

doc.add_paragraph()

# ── SECTION 5: HIGH LIQUIDITY SESSION WINDOWS ───────────────────────────────
add_heading(doc, "5. High Liquidity Session Windows (New)", level=1, color_hex="1F3864")

add_body(doc,
    "The system is designed to flourish during high-liquidity, directional market windows. "
    "The following five windows are the canonical active periods. Each bot family specifies "
    "its session mask from this list.")

make_table(doc,
    ["Window", "GMT Time", "Pairs", "Primary Strategy", "Priority"],
    [
        ["London Open", "07:00–09:00", "EUR, GBP, CHF and USD crosses", "ORB primary + Scalping", "Tier 1"],
        ["London–NY Overlap", "13:00–16:00", "All major pairs", "Scalping primary + ORB secondary", "Tier 1 — Premium"],
        ["NY Open / US Data", "13:00–15:30", "USD pairs, EURUSD, GBPUSD, USDJPY", "ORB secondary + Scalping", "Tier 1"],
        ["Tokyo Open", "00:00–02:00", "JPY pairs: USDJPY, EURJPY, GBPJPY. AUD active.", "Scalping — JPY pool", "Tier 2"],
        ["Sydney–Tokyo Transition", "22:00–01:00", "AUD, NZD, JPY pairs", "Scalping — AUD pool (wider spread filters required)", "Tier 2"],
    ],
    col_widths=[1.5, 1.2, 2.0, 2.0, 0.9]
)
add_body(doc, "Avoid: 16:00–22:00 GMT (post-London close to Sydney open). Low volume, widening spreads, "
         "choppy price action. This window is blocked for all new entries by session mask.")
add_note(doc, "The London–NY Overlap (Tier 1 Premium) is the highest-priority window and receives "
         "session-scoped Kelly modifications (see Section 9).")

doc.add_paragraph()

# ── SECTION 6: BOT FAMILY TEMPLATES ─────────────────────────────────────────
add_heading(doc, "6. Bot Family Templates (New)", level=1, color_hex="1F3864")

add_body(doc,
    "These templates define the specific rules for each strategy family. The AlphaForge agentic "
    "system MUST use these templates when generating EA code and validation workflows. When "
    "Development Department builds a bot, it selects a family template and adheres to the "
    "rules below. These are not suggestions — they are architectural constraints.")

add_heading(doc, "6.1 Scalping Family Template", level=2, color_hex="2E5496")
make_table(doc,
    ["Parameter", "Rule"],
    [
        ["Timeframes", "M1 (entry timing), M5 (near-term momentum), M15 (intraday structure), H1 (directional envelope)"],
        ["Holding period", "2–15 minutes. Time-based CET: if open > 2× average hold with no progress → close."],
        ["Direction", "Set by H1/H4 regime confidence via pool activation. EA only takes entries matching active pool."],
        ["Entry signal", "M1/M5 candle pattern + RVOL > 1.2 confirmation + VWAP proximity filter (from SVSS)"],
        ["SL — Layer 1 (EA)", "Hard SL = 1.5–2× ATR(M5). Set at order placement. Never removed."],
        ["TP — Layer 1 (EA)", "Hard TP = 1.2–1.5× SL distance. Fixed. EA manages this."],
        ["Layer 2 trigger", "H1 flip confirmed by Regime Persistence Timer (3× M5 bars). Move SL to break-even if 1R in profit."],
        ["Layer 3 trigger", "CHAOS only. Trust the EA for all normal regime cycling."],
        ["CET exits", "1) Fixed TP hit. 2) Fixed SL hit. 3) Momentum exhaustion: RVOL < 0.6 while open. 4) Time-based exit. 5) CHAOS signal."],
        ["Session masks", "London Open, London–NY Overlap, NY Open, Tokyo Open (JPY variants only)"],
        ["Spread quality", "SQS must be > 0.75 for entry. Auto-block 15 min pre/post high-impact news."],
        ["Trailing stop", "NOT used. Fixed TP/SL only. Regime-aware exit via CET framework instead."],
        ["R:R target", "1.2:1 minimum. Backtest rejection if below 1.1:1 net of fees."],
    ],
    col_widths=[1.8, 5.0]
)

add_heading(doc, "6.2 ORB Family Template", level=2, color_hex="2E5496")
make_table(doc,
    ["Parameter", "Rule"],
    [
        ["Timeframes", "M15 (range definition), H1 (breakout confirmation + regime alignment), H4 (macro direction)"],
        ["Holding period", "30 minutes to full session close. Force-close at session end (Islamic compliance / halal)."],
        ["Direction", "Set at session open by breakout direction. Fixed for the session. No intra-session pool switching."],
        ["Entry signal", "Price breaks session range high/low + RVOL > 1.5 surge + H1 regime aligned with breakout"],
        ["SL — Layer 1 (EA)", "Hard SL = below/above session range opposite boundary. Set at placement."],
        ["TP — Layer 1 (EA)", "Hard TP = 2.0–3.0× SL distance (ORB trades carry higher R:R target)"],
        ["Layer 2 trigger", "H1+H4 both flip against trade, confirmed by Regime Persistence Timer (2× H1 bars) → tighten stop. NOT close."],
        ["Layer 3 trigger", "CHAOS only. Tighten stop (not immediate close) unless kill switch."],
        ["CET exits", "1) Fixed TP hit. 2) Fixed SL hit. 3) Regime: BREAKOUT→RANGE confirmed → tighten stop. 4) Session end forced close."],
        ["Session masks", "London Open primary (07:00–10:00 GMT). NY Open secondary (13:00–16:00 GMT)."],
        ["Spread quality", "SQS > 0.80 for ORB entries (stricter than scalping — ORB entries have wider stops, spread impact is larger)"],
        ["Trailing stop", "Regime-aware tightening only — NOT distance-based trailing. Stop moves when Sentinel shifts."],
        ["Capital routing", "Sniper account (Exness Raw / IC Markets). Kelly naturally produces larger fraction due to higher R:R."],
        ["False Breakout variant", "Separate pool (Section 4.2). Distinct backtest history. Activated only on confirmed reversal signal."],
    ],
    col_widths=[1.8, 5.0]
)

doc.add_paragraph()

# ── SECTION 7: SHARED VOLUME SESSION SERVICE (SVSS) ─────────────────────────
add_heading(doc, "7. Shared Volume Session Service — SVSS (New)", level=1, color_hex="1F3864")

add_body(doc,
    "A critical gap in the current architecture. Volume Profile, VWAP, and RVOL calculations "
    "must be shared services — not per-bot computations. 200 bots independently computing VWAP "
    "on the same instrument is 200 redundant calculations. One service, one cache, 200 consumers.")

add_heading(doc, "7.1 Architecture", level=2, color_hex="2E5496")
add_body(doc, "SVSS runs on Cloudzy, alongside the live tick feed. For each subscribed instrument:")
add_bullet(doc, "VWAP: Volume-weighted average price. Running sum reset at each session open. "
           "Formula: Σ(price × volume) / Σ(volume). Updated on every tick.")
add_bullet(doc, "Volume Profile: Price-level histogram bucketed to 0.1-pip resolution. "
           "Identifies Point of Control (POC), Value Area High (VAH), Value Area Low (VAL).")
add_bullet(doc, "RVOL (Relative Volume): current_bar_volume / avg_volume_at_this_time_of_day "
           "(rolling 20-session average). Updated on bar close.")
add_bullet(doc, "MFI (Money Flow Index): 14-period. Combines price direction + volume to detect divergences. "
           "Overbought > 80 = momentum exhaustion signal for CET exits.")

add_heading(doc, "7.2 Data Source", level=2, color_hex="2E5496")
add_body(doc,
    "MT5 via IC Markets provides real contract volume (lot-based, not tick count) accessible "
    "via iVolume() in MQL5 and through the ZMQ tick feed. SVSS consumes the ZMQ tick feed "
    "directly for its calculations, independent of the MT5 built-in indicators. MT5 native "
    "volume indicators may be used for EA-level entry filters; SVSS provides the shared "
    "intelligence layer for the Python strategy router.")

add_heading(doc, "7.3 Publishing Pattern", level=2, color_hex="2E5496")
add_body(doc,
    "SVSS publishes all calculations to Redis cache on each tick update and bar close. "
    "All bots read from Redis — no direct SVSS API calls. "
    "Key pattern: svss:{symbol}:vwap, svss:{symbol}:rvol, svss:{symbol}:poc, etc.")

add_heading(doc, "7.4 SVSS → Correlation Sensor Connection", level=2, color_hex="2E5496")
add_body(doc,
    "When RVOL spikes simultaneously across multiple instruments (e.g., EURUSD, GBPUSD, "
    "USDJPY all show RVOL > 2.0 at the same moment), this is a systemic liquidity event — "
    "a macro news release or central bank statement. The Correlation Sensor's max_eigenvalue "
    "will simultaneously spike. These two signals together trigger CHAOS classification in "
    "the Sentinel. This is the early warning chain: SVSS → CorrelationSensor → Sentinel → "
    "Layer 3 active.")

doc.add_paragraph()

# ── SECTION 8: DATA PIPELINE LATENCY BUDGET ─────────────────────────────────
add_heading(doc, "8. Data Pipeline Latency Budget (New)", level=1, color_hex="1F3864")

add_body(doc,
    "All components below run on Cloudzy except where noted. Expensive recalculations "
    "(Sentinel, Kelly, Correlation) are event-driven, not per-tick.")

make_table(doc,
    ["Stage", "Update Trigger", "Target Latency", "Note"],
    [
        ["IC Markets WebSocket → Cloudzy", "Every tick", "1–3ms", "LD4 to LD4 co-location target"],
        ["OHLCV aggregator (tick → candle)", "Every tick", "< 1ms", "In-memory running sums"],
        ["SVSS update (VWAP, RVOL, MFI)", "Every tick / bar close", "< 1ms", "Running calculation, not batch"],
        ["MultiTimeframeSentinel regime", "Bar close only (M1 = 60s, M5 = 5min)", "5–15ms (on bar close), 0ms (cached)", "Never recalculates mid-bar"],
        ["CorrelationSensor (RMT)", "Every 5 minutes", "10–20ms (calc), 0ms (cached)", "Contabo or Cloudzy"],
        ["Kelly Engine recalculation", "Regime change event only", "< 5ms (cache read between changes)", "NOT per-tick — regime-event-driven"],
        ["Governor per-bot evaluation", "Per incoming signal", "< 5ms", "Redis atomic lookup + Kelly cache"],
        ["Commander routing matrix", "Per approved signal", "< 2ms", "Pre-computed O(1) lookup"],
        ["ZMQ order dispatch + broker RT", "Per approved trade", "5–20ms", "Cloudzy LD4 → IC Markets LD4"],
        ["Total signal-to-order", "—", "12–35ms", "Within 5–20ms execution tier target"],
    ],
    col_widths=[2.1, 1.8, 1.5, 2.3]
)

doc.add_paragraph()

# ── SECTION 9: KILL SWITCH INTEGRATION ──────────────────────────────────────
add_heading(doc, "9. Kill Switch Integration with New Components (Addendum)", level=1, color_hex="1F3864")

add_body(doc,
    "The existing ProgressiveKillSwitch and SmartKillSwitch are confirmed. This section "
    "clarifies how new components added in this session interact with the kill switch:")

add_bullet(doc, "SVSS and CorrelationSensor: Continue running during all kill switch tiers. "
           "They are monitoring infrastructure, not trading infrastructure.")
add_bullet(doc, "Layer 2 Position Monitor: On Tier 1 (soft stop) — tighten all open stops to "
           "break-even. On Tier 2 — move stops to capture partial profit. On Tier 3 — close all. "
           "The Position Monitor is the delivery mechanism for kill switch commands to open positions.")
add_bullet(doc, "Regime-Conditional Pools: On kill switch activation — no new pool activations "
           "regardless of regime signal. The Commander's auction is suspended.")
add_bullet(doc, "Session-Scoped Kelly Modifiers: Kill switch overrides all Kelly multipliers. "
           "No house money expansion during kill switch state.")
add_bullet(doc, "Spread Quality Score: During kill switch recovery (Tier 1 lifted) — SQS threshold "
           "temporarily raised to 0.85 (from 0.75) for 30 minutes to ensure clean re-entry conditions.")

doc.add_paragraph()

# ── SECTION 10: CORRELATION SENSOR → GOVERNOR INTEGRATION ───────────────────
add_heading(doc, "10. Correlation Sensor → Governor Integration (New)", level=1, color_hex="1F3864")

add_body(doc,
    "The CorrelationSensor (src/risk/physics/correlation_sensor.py) uses Random Matrix Theory "
    "with Marchenko-Pastur distribution to detect genuine systemic correlation versus noise. "
    "It exists but is NOT yet wired into the Governor. This section specifies the integration.")

add_heading(doc, "10.1 Marginal Correlation Evaluation", level=2, color_hex="2E5496")
add_body(doc,
    "Before the Governor approves a new trade, it calculates the marginal correlation impact "
    "of adding that position to the existing portfolio. The formula:")

add_code(doc, "# For existing portfolio positions i with weights w_i and correlation matrix C:")
add_code(doc, "# Pairwise correlation penalty between position i and proposed position j:")
add_code(doc, "")
add_code(doc, "correlation_penalty(i,j) = max(0, C_ij - threshold) / (1 - threshold)")
add_code(doc, "adjusted_kelly_j = base_kelly_j × (1 - correlation_penalty(i,j))")
add_code(doc, "")
add_code(doc, "# threshold = 0.5 (configurable in risk/config.py)")
add_code(doc, "# Example: EURUSD and GBPUSD, C_ij = 0.80:")
add_code(doc, "# penalty = (0.80 - 0.50) / (1 - 0.50) = 0.60")
add_code(doc, "# GBPUSD position gets 40% of its base Kelly size if EURUSD already open")

add_heading(doc, "10.2 Does Timeframe Matter for Correlation?", level=2, color_hex="2E5496")
add_body(doc,
    "Yes. Correlation between instruments varies by timeframe. On M1, EUR/GBP correlation "
    "may be 0.60 (noise-dominant, short-horizon). On H1, it may be 0.85 (genuine macro-driven). "
    "The CorrelationSensor should be run separately for each timeframe the bot family uses: "
    "Scalping bots use M5 correlation matrix. ORB bots use H1 correlation matrix. "
    "Governor selects the correct matrix based on the bot's family template timeframe.")

add_heading(doc, "10.3 Update Frequency and Caching", level=2, color_hex="2E5496")
add_body(doc,
    "CorrelationSensor recalculates every 5 minutes on bar close data (not per-tick). "
    "Result cached in Redis: risk:correlation:M5 and risk:correlation:H1. Governor reads "
    "from cache — zero computation per trade evaluation. The 5-minute update window is "
    "sufficient because genuine correlation regime shifts occur over minutes to hours, not seconds.")

add_heading(doc, "10.4 Portfolio Variance Constraint", level=2, color_hex="2E5496")
add_body(doc,
    "If adding a proposed trade increases the portfolio's correlation-adjusted total exposure "
    "above the 5% concurrent cap, the Governor rejects it — same as any other cap breach. "
    "The correlation adjustment means that two highly correlated positions (EURUSD + GBPUSD "
    "both long) are treated as ~1.4 positions of exposure, not 2 independent 0.5% positions.")

doc.add_paragraph()

# ── SECTION 11: RVOL-WEIGHTED POSITION SIZING ───────────────────────────────
add_heading(doc, "11. RVOL-Weighted Position Sizing (New)", level=1, color_hex="1F3864")

add_body(doc,
    "RVOL (Relative Volume) = current bar volume / average volume at this time of day "
    "(rolling 20-session baseline). RVOL measures execution quality, not directional signal.")

add_body(doc,
    "After the Kelly engine produces its base position size, the Governor applies an RVOL "
    "execution quality multiplier:")

add_code(doc, "final_size = kelly_size × clamp(RVOL, 0.5, 1.5)")
add_code(doc, "")
add_code(doc, "RVOL ≥ 2.0  → trade at 1.5× Kelly (cap — don't over-size on extreme volume)")
add_code(doc, "RVOL = 1.0  → trade at 1.0× Kelly (normal conditions)")
add_code(doc, "RVOL = 0.6  → trade at 0.6× Kelly (thin market — reduce for slippage risk)")
add_code(doc, "RVOL < 0.5  → entry blocked (insufficient liquidity for scalping)")

add_body(doc,
    "Rationale: This is not directional speculation. Higher RVOL = better execution quality "
    "(tighter effective spread, less slippage, cleaner fills). The edge is more realisable "
    "in high-volume conditions. RVOL is provided by SVSS — no additional calculation required.")

doc.add_paragraph()

# ── SECTION 12: SESSION-SCOPED KELLY MODIFIERS ──────────────────────────────
add_heading(doc, "12. Session-Scoped Kelly Modifiers (New)", level=1, color_hex="1F3864")

add_body(doc,
    "A new architectural concept not present in any prior document. Standard house money "
    "thresholds apply globally. Session-scoped modifiers apply adjustments during specific "
    "windows without changing the baseline risk framework.")

add_heading(doc, "12.1 London–NY Overlap: Early House Money Threshold", level=2, color_hex="2E5496")
add_body(doc,
    "The London–NY Overlap (13:00–16:00 GMT) is the system's premium trading window — "
    "highest liquidity, tightest spreads, most reliable signals. During this window, the "
    "house money activation threshold is lowered:")

make_table(doc,
    ["Condition", "Normal Sessions", "London–NY Overlap"],
    [
        ["House money activates when", "Equity > 1.5× initial deposit", "Equity > 1.2× initial deposit"],
        ["Kelly fraction multiplier at activation", "0.35 (vs 0.25 base)", "0.35 (same — earlier entry only)"],
        ["Rationale", "Standard profit buffer trigger", "Better conditions justify earlier scaling"],
    ],
    col_widths=[2.5, 2.1, 2.1]
)

add_heading(doc, "12.2 Reverse House Money Effect (London–NY Overlap Only)", level=2, color_hex="2E5496")
add_body(doc,
    "If the London–NY Overlap session produces consecutive losses, a session-scoped reverse "
    "house money modifier reduces position sizing progressively. This protects against handing "
    "back a profitable week in one bad premium session.")

make_table(doc,
    ["Trigger", "Kelly Modifier Applied", "State"],
    [
        ["2 consecutive losses in this session", "Kelly × 0.70 (−30%)", "Cautious mode"],
        ["4 consecutive losses in this session", "Kelly × 0.50 (−50%)", "Defensive mode"],
        ["Daily drawdown limit hit", "Trading suspended (existing 3/5/7 rule)", "Existing framework"],
    ],
    col_widths=[2.5, 2.0, 1.7]
)
add_body(doc,
    "All session-scoped modifiers reset to baseline when the session window closes (16:00 GMT). "
    "Next session starts with clean Kelly state. The system does NOT require operator permission "
    "to apply these modifiers — they activate and deactivate automatically based on performance "
    "metrics and the session window. Automation is the design intent.")

doc.add_paragraph()

# ── SECTION 13: SPREAD QUALITY SCORE ────────────────────────────────────────
add_heading(doc, "13. Spread Quality Score (SQS) System (New)", level=1, color_hex="1F3864")

add_heading(doc, "13.1 Architecture", level=2, color_hex="2E5496")
add_body(doc,
    "The system maintains a historical spread table per instrument capturing the average spread "
    "at each 5-minute interval of the trading day across the last 30 sessions. This produces a "
    "'normal spread profile' curve for each instrument.")

add_code(doc, "SQS = historical_avg_spread_at_this_time / current_live_spread")
add_code(doc, "")
add_code(doc, "SQS = 1.0  → normal spread conditions")
add_code(doc, "SQS = 0.7  → spread 30% wider than normal (degraded execution quality)")
add_code(doc, "SQS = 1.3  → spread narrower than usual (optimal conditions)")

add_heading(doc, "13.2 Entry Filters", level=2, color_hex="2E5496")
make_table(doc,
    ["Strategy Family", "SQS Threshold for Entry", "Rationale"],
    [
        ["Scalping", "SQS > 0.75", "Tight stops make spread cost significant relative to R"],
        ["ORB", "SQS > 0.80", "Stricter — ORB has wider stops but longer hold; spread at entry matters"],
        ["Both (hard block)", "SQS < 0.50", "Spread doubled or more — no entry regardless of signal"],
    ],
    col_widths=[1.7, 1.8, 3.2]
)

add_heading(doc, "13.3 Economic Calendar Integration", level=2, color_hex="2E5496")
add_body(doc,
    "High-impact events (FOMC, NFP, CPI, ECB, BOE decisions) are loaded weekly from the "
    "economic calendar. The system applies:")
add_bullet(doc, "Pre-event blackout: −15 minutes before release. No new entries for affected pairs.")
add_bullet(doc, "Post-event cooldown: +15 minutes after release OR until SQS returns to > 0.75, "
           "whichever is later.")
add_bullet(doc, "Affected pair detection: calendar event metadata maps to instrument groups "
           "(e.g., US CPI → EURUSD, GBPUSD, USDJPY blocked).")
add_note(doc, "Economic calendar infrastructure already exists in the codebase (calendar_governor.py). "
         "The SQS system connects to it as a supplementary filter, not a replacement.")

doc.add_paragraph()

# ── SECTION 14: HMM — 3-DAY FEEDBACK LAG & WALK-FORWARD CALIBRATOR ──────────
add_heading(doc, "14. HMM: 3-Day Feedback Lag & Walk-Forward Calibrator (Addendum)", level=1, color_hex="1F3864")

add_body(doc,
    "The HMM is in shadow mode (ISING_ONLY → HMM_SHADOW) for 6–8 months while IC Markets live "
    "data accumulates. Two design decisions are added here.")

add_heading(doc, "14.1 3-Day Feedback Lag", level=2, color_hex="2E5496")
add_body(doc,
    "When a live trade closes, its outcome data does NOT immediately enter the HMM training pool. "
    "A mandatory 3-day delay is enforced before the trade record is eligible for training input. "
    "Purpose: prevents the system from training on its own recent live decisions in a feedback loop "
    "(a form of data leakage that would degrade the model over time). After 3 days, the data is "
    "treated as historical and enters the normal distillation pipeline.")
add_note(doc, "This applies to both HMM retraining and AlphaForge Workflow 2 improvement loop. "
         "The Dataset Copy Firewall (from prior planning document) enforces the time boundary — "
         "this 3-day lag is an additional rule applied on top of that existing mechanism.")

add_heading(doc, "14.2 Walk-Forward Window Calibrator", level=2, color_hex="2E5496")
add_body(doc,
    "The current architecture specifies rolling 1-month windows for WFA on scalping strategies. "
    "This should be dynamic, calibrated to the HMM's observed regime transition frequency:")
add_code(doc, "wfa_window = f(avg_regime_transition_interval)")
add_code(doc, "")
add_code(doc, "If regimes change every 3 weeks on average:")
add_code(doc, "  → 1-month window is appropriate (covers ~1.3 regime transitions)")
add_code(doc, "If regimes change every 6 weeks on average:")
add_code(doc, "  → 1-month window is too short; use 6-week window")
add_code(doc, "If regimes change every 10 days:")
add_code(doc, "  → 1-month window is too long; use 2-week window")
add_body(doc,
    "The HMM on Contabo tracks regime transition timestamps. The AlphaForge Research sub-agent "
    "queries the average transition interval before setting up a WFA run and passes the calibrated "
    "window size to the backtest engine. This is an agent tool call, not a hardcoded constant.")

doc.add_paragraph()

# ── SECTION 15: PORTFOLIO MATH — EXPECTED VALUE AT SCALE ────────────────────
add_heading(doc, "15. Portfolio Math — Expected Value at Scale", level=1, color_hex="1F3864")

add_heading(doc, "15.1 Base Parameters", level=2, color_hex="2E5496")
make_table(doc,
    ["Parameter", "Value", "Notes"],
    [
        ["Win rate (W)", "52%", "Phase 1 target. Agentic loop maintains ≥ 52% over time."],
        ["R:R ratio (scalping)", "1.2:1", "Gross. Net after spread/commission closer to 1.05:1."],
        ["Per-trade risk (R)", "0.5% of equity", "Base Kelly fraction × quarter-Kelly discipline"],
        ["Concurrent exposure cap", "5% of equity", "Governor hard limit"],
        ["Max concurrent open positions", "10 (at any risk %) ", "(5% cap) / (0.5% per trade) = 10"],
        ["Average holding time", "8 minutes", "Scalping. ORB = 30–180 minutes."],
        ["Commission cost per trade", "~0.035% equity", "IC Markets raw at growing lot sizes"],
    ],
    col_widths=[2.2, 1.8, 3.2]
)

add_heading(doc, "15.2 Expected Value Calculation", level=2, color_hex="2E5496")
add_code(doc, "EV_gross = (W × avg_win_R) - (L × avg_loss_R)")
add_code(doc, "EV_gross = (0.52 × 1.2R) - (0.48 × 1.0R) = 0.624R - 0.480R = +0.144R per trade")
add_code(doc, "")
add_code(doc, "EV_net   = EV_gross - commission_per_trade")
add_code(doc, "EV_net   ≈ 0.144R - 0.050R = +0.094R per trade (conservative fee estimate)")
add_code(doc, "As % of equity: 0.094 × 0.5% = +0.047% net per trade")

add_heading(doc, "15.3 Scale Scenarios", level=2, color_hex="2E5496")
make_table(doc,
    ["Stage", "Account", "Trades/Day", "Daily EV (net)", "Daily Return", "Monthly (~22 days)", "12-Month Growth"],
    [
        ["Start", "$200", "30", "$0.28", "0.14%", "~3%", "~$236"],
        ["Early Growth", "$500", "50", "$1.18", "0.24%", "~5.3%", "~$774"],
        ["Intermediate", "$2,000", "150", "$14.10", "0.71%", "~16.7%", "large"],
        ["Full Scale (50 bots)", "$5,000", "300", "$70.50", "1.4%", "~36%", "compounding"],
        ["Full Scale (session-adjusted, 5 sessions)", "$5,000", "450–600", "$105–$141", "2.1–2.8%", "~57–81%", "—"],
    ],
    col_widths=[1.5, 0.9, 1.0, 1.4, 1.1, 1.4, 1.4]
)

add_heading(doc, "15.4 The Agentic Improvement Loop as the Compounding Multiplier", level=2, color_hex="2E5496")
add_body(doc,
    "The math above assumes a static 52% win rate. In practice, strategies decay over time as "
    "market conditions shift. The Decline and Recovery Loop (prior planning document) is what "
    "maintains the portfolio-level win rate at or above threshold. A static system at 52% may "
    "decay to 49% in 3 months. An agentic system that continuously diagnoses underperformers, "
    "quarantines them, and improves or replaces them should hold at ≥52% as a maintained minimum "
    "and potentially improve to 54–55% on optimised quarters. The math of compound growth at "
    "maintained 52%+ is substantial — the agentic loop is not a nice-to-have, it is the "
    "mechanism that makes the long-term EV projections achievable.")

add_heading(doc, "15.5 Daily Net EV in Workflows", level=2, color_hex="2E5496")
add_body(doc,
    "The daily net expected value figure (per-trade EV × daily trade count) should be tracked "
    "as a workflow metric in the Bot Variant Lifecycle Report system. The Live Monitor sub-agent "
    "compares actual daily EV against the expected EV from backtest at each lifecycle stage. "
    "A delta of > −20% for 5 consecutive trading days triggers the Decline and Recovery Loop.")

doc.add_paragraph()

# ── SECTION 16: DEFERRED AND REMOVED ITEMS ───────────────────────────────────
add_heading(doc, "16. Deferred and Explicitly Removed Items", level=1, color_hex="1F3864")

add_heading(doc, "16.1 Removed", level=2, color_hex="2E5496")
make_table(doc,
    ["Item", "Decision", "Reason"],
    [
        ["Intraday P&L Lock Mode", "REMOVED", "Limits the system unnecessarily. Session mask + drawdown limits are the only runtime constraints."],
        ["Alpaca API integration", "REMOVED from scope", "US equities broker. Irrelevant to forex/MT5 stack. Revisit only if equity strategies are ever added."],
        ["7% London-NY overlap exposure cap", "REPLACED", "Replaced by session-scoped Kelly modifiers (Section 12). No raw exposure cap increase."],
    ],
    col_widths=[2.0, 1.4, 3.3]
)

add_heading(doc, "16.2 Deferred to Phase 2", level=2, color_hex="2E5496")
make_table(doc,
    ["Item", "Phase", "Condition for Activation"],
    [
        ["Variant Genetic Fingerprint", "Phase 2", "After agentic system is fully operational and tested. Alpha Forge producing live variants."],
        ["Walk-Forward Window Auto-Calibration (full)", "Phase 1 → Phase 2", "Requires HMM to have sufficient transition data (6+ months). Phase 1: use fixed 1-month window."],
        ["Cross-Market Regime Score (DXY, Gold, US10Y)", "Phase 2", "Requires secondary data source for DXY/US10Y. Deferred per prior planning document."],
        ["HMM promotion to HMM_HYBRID", "6–8 months", "Shadow mode accumulates IC Markets data. Promoted when prediction accuracy > 70%."],
    ],
    col_widths=[2.2, 0.9, 3.6]
)

doc.add_paragraph()

# ── SECTION 17: OPEN QUESTIONS ───────────────────────────────────────────────
add_heading(doc, "17. Open Questions", level=1, color_hex="1F3864")

make_table(doc,
    ["Question", "Current Thinking", "Decision Point"],
    [
        ["Exness → IC Markets migration timing",
         "Migrate when account is stable and Eversend setup is complete. No code changes required.",
         "When infrastructure ready"],
        ["ORB False Breakout: separate account or same Sniper account?",
         "Same Sniper account, separate magic number range (20000–20999). Separate backtest history.",
         "Implementation phase"],
        ["Scalping variant count: how many variants reach 45 active bots?",
         "One base strategy can produce 8 variants (M1/M5 × long/short × 2 param sweeps). "
         "6 base strategies × 8 variants = 48 scalping bots. Achievable.",
         "Alpha Forge Workflow 1"],
        ["Session-scoped Kelly modifier persistence: carry to next day if session ends mid-drawdown?",
         "No. Modifiers reset at session close. Each session starts fresh.",
         "Confirmed this session"],
        ["RVOL baseline window (20 sessions): is this sufficient for low-liquidity sessions?",
         "Tokyo/Sydney may need 30-session baseline due to higher day-to-day variance. Configurable.",
         "Implementation phase"],
    ],
    col_widths=[2.0, 3.0, 1.7]
)

doc.add_paragraph()

# ── FOOTER ───────────────────────────────────────────────────────────────────
add_heading(doc, "Document References", level=1, color_hex="1F3864")
add_bullet(doc, "QuantMindX_Planning_Document_March2026.docx — Primary planning document (prior session)")
add_bullet(doc, "_bmad-output/planning-artifacts/prd.md — Product Requirements (80 FRs, 52 journeys)")
add_bullet(doc, "_bmad-output/planning-artifacts/architecture.md — Architecture Decision Document")
add_bullet(doc, "_bmad-output/planning-artifacts/epics.md — Epic and Story Breakdown")
add_bullet(doc, "_bmad-output/planning-artifacts/ux-design-specification.md — UX/Frosted Terminal spec")
add_bullet(doc, "src/risk/physics/correlation_sensor.py — Existing CorrelationSensor (RMT implementation)")
add_bullet(doc, "src/router/routing_matrix.py — Existing RoutingMatrix (Machine Gun / Sniper accounts)")
add_bullet(doc, "src/position_sizing/enhanced_kelly.py — Existing EnhancedKellyCalculator")
add_bullet(doc, "src/router/multi_timeframe_sentinel.py — Existing MultiTimeframeSentinel")

doc.add_paragraph()
final_p = doc.add_paragraph()
final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = final_p.add_run("QuantMindX | Planning Addendum — Session 2 | March 2026")
fr.font.size = Pt(9)
fr.font.italic = True
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── SAVE ─────────────────────────────────────────────────────────────────────
out_path = "/home/mubarkahimself/Desktop/QUANTMINDX/claude-desktop-workfolder/QuantMindX_Planning_Addendum_Session2_March2026.docx"
doc.save(out_path)
print(f"Saved: {out_path}")

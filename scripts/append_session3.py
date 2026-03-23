"""
QuantMindX Planning Addendum — Session 3 Append Script
Appends all new architectural decisions from Session 3 (March 2026) to the existing
Session 2 addendum. Does NOT delete any existing content.

Run from any directory:
    python3 scripts/append_session3.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ADDENDUM_PATH = "/home/mubarkahimself/Desktop/QUANTMINDX/claude-desktop-workfolder/QuantMindX_Planning_Addendum_Session2_March2026.docx"

# ─── Colour palette (matches Session 2 document style) ───────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # section heading backgrounds
MID_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # sub-heading text
LIGHT_BLUE  = RGBColor(0xD5, 0xE8, 0xF0)   # table header fill
PALE_GREY   = RGBColor(0xF2, 0xF2, 0xF2)   # table row alternating fill
ORANGE      = RGBColor(0xC5, 0x5A, 0x11)   # warning / note highlight
GREEN_DARK  = RGBColor(0x37, 0x5A, 0x32)   # confirmed / positive
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)


# ─── Helper: set paragraph shading ───────────────────────────────────────────
def shade_paragraph(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ─── Helper: add section heading (dark blue bar + white bold text) ────────────
def add_section_heading(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    shade_paragraph(p, "1F497D")
    run = p.add_run(f"  {number}. {title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    run.font.name = "Calibri"


# ─── Helper: add sub-heading (mid blue text) ─────────────────────────────────
def add_sub_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = MID_BLUE
    run.font.name = "Calibri"


# ─── Helper: add body paragraph ──────────────────────────────────────────────
def add_body(doc, text: str, bold_prefix: str = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"


# ─── Helper: add bullet ──────────────────────────────────────────────────────
def add_bullet(doc, text: str, level: int = 0, bold_prefix: str = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"


# ─── Helper: add labelled field (bold label + body text) ─────────────────────
def add_labeled(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"


# ─── Helper: add note box (orange background) ────────────────────────────────
def add_note(doc, text: str, label: str = "NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.25)
    shade_paragraph(p, "FFF2CC")
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    r.font.color.rgb = ORANGE
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = "Calibri"


# ─── Helper: add code block (grey background, monospace) ─────────────────────
def add_code(doc, text: str):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.3)
        shade_paragraph(p, "F2F2F2")
        r = p.add_run(line if line else " ")
        r.font.size = Pt(9)
        r.font.name = "Courier New"


# ─── Helper: make a simple table ─────────────────────────────────────────────
def make_table(doc, headers: list, rows: list, col_widths: list = None):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, "2E75B6")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        fill = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(str(cell_text))
            r.font.size = Pt(10)
            r.font.name = "Calibri"

    # Set column widths
    if col_widths:
        for row in table.rows:
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Inches(width)

    doc.add_paragraph()  # spacer after table


# ─── Helper: divider line ────────────────────────────────────────────────────
def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    shade_paragraph(p, "2E75B6")
    r = p.add_run(" ")
    r.font.size = Pt(3)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN — load and append
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document(ADDENDUM_PATH)

# ── SESSION 3 SEPARATOR ──────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(6)
shade_paragraph(p, "1F497D")
r = p.add_run("  SESSION 3 ADDENDUM — March 2026 (Continued)")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = WHITE
r.font.name = "Calibri"

add_body(doc,
    "The following sections append new architectural decisions, frameworks, and clarifications "
    "arising from Session 3. No content from Sessions 1 or 2 has been removed. A coding agent "
    "reading this document should treat Session 3 content as refinements and extensions to "
    "everything above. Where a Session 3 decision supersedes a Session 2 decision, the Session 3 "
    "version takes precedence. All terminology is defined at first use in each section.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-1  CORRECTED PORTFOLIO MATHEMATICS
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-1", "Corrected Portfolio Mathematics")

add_sub_heading(doc, "S3-1.1  Confirmed Trading Parameters")
add_labeled(doc, "Starting equity",         "$200 USD")
add_labeled(doc, "Risk per trade",          "2% of current equity (Kelly hard cap). At $200 = $4 per trade. NOT $2 — the $2 figure was based on a 1% assumption from an earlier draft. 2% is the EnhancedKellyConfig.max_risk_pct default.")
add_labeled(doc, "R:R ratio target",        "1:2 minimum (win = 2× the risk). 1:3 optimal (win = 3× the risk). The 1.2:1 figure used in earlier documents is incorrect for this system — that was sourced from generic BMAD defaults.")
add_labeled(doc, "Fees",                    "Already baked into Kelly via fee-adjusted avg_win / avg_loss calculation in EnhancedKellyCalculator.calculate(). Do NOT deduct fees separately from EV calculations.")
add_labeled(doc, "Win rate baseline",       "52% across validated strategy families. 50% used for break-even analysis.")

add_sub_heading(doc, "S3-1.2  Per-Trade EV at Confirmed Parameters")
add_body(doc,
    "At 2% risk per trade on a $200 account, risk per trade = $4. "
    "Win = $8 (1:2 R:R). Loss = $4.")

make_table(doc,
    ["Scenario", "Win$", "Loss$", "EV per Trade", "As % of Equity"],
    [
        ["1:2 R:R, 50% WR",  "$8.00", "$4.00", "+$2.00", "+1.00%"],
        ["1:2 R:R, 52% WR",  "$8.00", "$4.00", "+$2.24", "+1.12%"],
        ["1:3 R:R, 50% WR",  "$12.00","$4.00", "+$4.00", "+2.00%"],
        ["1:3 R:R, 52% WR",  "$12.00","$4.00", "+$4.32", "+2.16%"],
    ],
    col_widths=[2.2, 1.1, 1.1, 1.4, 1.4]
)

add_sub_heading(doc, "S3-1.3  Daily EV at 30 Trades (Early Phase)")
make_table(doc,
    ["Scenario", "EV/Trade", "× 30 trades", "Daily Net", "Day-1 Equity"],
    [
        ["1:2 R:R, 50% WR", "+$2.00", "30", "+$60.00", "$260"],
        ["1:2 R:R, 52% WR", "+$2.24", "30", "+$67.20", "$267"],
        ["1:3 R:R, 50% WR", "+$4.00", "30", "+$120.00","$320"],
        ["1:3 R:R, 52% WR", "+$4.32", "30", "+$129.60","$330"],
    ],
    col_widths=[2.2, 1.2, 1.1, 1.3, 1.4]
)

add_sub_heading(doc, "S3-1.4  Daily EV at 80 Trades (Full Phase, Multiple Bots)")
add_body(doc,
    "At full operational scale, the concurrent exposure cap (5% max = 5 bots at 2% each) "
    "limits simultaneous open trades to 5. However, since scalping trades close in under 20 "
    "minutes on average, the system cycles through multiple rounds of 5-concurrent bots per "
    "session. Using Little's Law (L = lambda × W): with L=5 slots and W=15 min avg hold time, "
    "throughput = 1 trade per 3 minutes = 20 trades/hour. A 4-hour active window = 80 trades/day "
    "without violating the 5% concurrent cap.")
make_table(doc,
    ["Scenario", "EV/Trade", "× 80 trades", "Daily Net", "Daily Return %"],
    [
        ["1:2 R:R, 52% WR", "+$2.24", "80", "+$179.20", "+89.6%"],
        ["1:3 R:R, 52% WR", "+$4.32", "80", "+$345.60", "+172.8%"],
    ],
    col_widths=[2.2, 1.2, 1.2, 1.4, 1.5]
)
add_note(doc,
    "These are theoretical expected values assuming consistent 52% WR and 80 trades firing every day. "
    "In practice, regime filters, SQS gates, and session masks will reduce active trade count on "
    "low-quality days. These numbers represent the system ceiling, not the daily guarantee. "
    "The concurrent exposure cap (5% portfolio max) is a hard limit that does not change regardless "
    "of trade count. Risk is controlled. Throughput is high because trades are short.")

add_sub_heading(doc, "S3-1.5  Compound Growth Projections (1:2 R:R, 52% WR, 80 trades/day)")
add_body(doc,
    "Because risk is always 2% of CURRENT equity, dollar gains grow as equity grows. "
    "The daily return % stays constant; the absolute dollar amount compounds.")
make_table(doc,
    ["Day", "Equity", "Daily $ Gain (approx)", "Notes"],
    [
        ["1",    "$379",      "$179",  "Day 1 at 80-trade full deployment"],
        ["5",    "$2,700",    "$1,279","After 5 trading days"],
        ["10",   "$19,300",   "$9,100","Compounding effect visible"],
        ["22",   "Theoretical ceiling — variance dominates beyond Day 10", "", ""],
    ],
    col_widths=[0.6, 1.5, 2.0, 2.5]
)
add_note(doc,
    "IMPORTANT FOR CODING AGENT: These projections are theoretical math, not a trading plan. "
    "The system is designed to approach these numbers asymptotically as the bot library grows "
    "and strategies are validated. Starting phase (Months 1-2) will see 8-15 trades/day, not 80. "
    "The numbers justify the architecture — they do not define the launch target.",
    label="CODING AGENT NOTE")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-2  SL/TP HYBRID EXIT LOGIC (OPTION B + C)
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-2", "SL/TP Hybrid Exit Logic — Option B+C")

add_sub_heading(doc, "S3-2.1  Confirmed Exit Strategy")
add_body(doc,
    "The confirmed exit strategy for scalping bots is a hybrid of Option B (breakeven move) "
    "and Option C (partial close + runner). Option D (dynamic CET) is reserved for Phase 2 "
    "bots with 6+ months of validated live performance. This section defines the exact mechanics "
    "so the EA template can implement them correctly.")

add_sub_heading(doc, "S3-2.2  Scalping Exit Mechanics (1:2 R:R Target)")
add_labeled(doc, "Entry",            "Set SL at X pips below entry. Set initial TP at 2X pips (1:2 target).")
add_labeled(doc, "At +0.5× SL pips","Move SL to breakeven. The full-loss scenario is now eliminated. Worst outcome = 0R.")
add_labeled(doc, "At +1× SL pips (1:1 level)", "Close 30% of the position. Lock that profit. Remaining 70% continues running with SL at breakeven.")
add_labeled(doc, "Runner TP",        "20 pips (1:2) for the 70% runner. If RVOL > 1.4 AND spread still clean AND regime still confirmed, EA may extend TP to 30 pips (1:3). Otherwise close at 1:2.")
add_labeled(doc, "Average effective R:R", "1.7R to 2.4R per winning trade depending on runner outcome. Better than static 1:2.")

add_sub_heading(doc, "S3-2.3  ORB Exit Mechanics (1:3 R:R Target)")
add_labeled(doc, "TP1", "1:1 level — close 50% of the position.")
add_labeled(doc, "SL move", "Move SL to breakeven after TP1 hit.")
add_labeled(doc, "Runner", "50% position continues to 1:3 TP. ORB moves have institutional follow-through so the runner gets more room.")

add_sub_heading(doc, "S3-2.4  The Four Stuck-in-Middle Exit Rules")
add_body(doc, "A trade that is open between breakeven SL and TP — not progressing — consumes a queue slot. "
    "These four rules define what happens. All four are implemented in the EA template.")
add_bullet(doc, "MAXIMUM HOLD TIME: Momentum scalpers 20 min max. Liquidity grab scalpers 15 min max. "
    "Session open scalpers 30 min max. ORB bots 4 hours max. If not resolved within hold time, close at market.", bold_prefix="Rule 1")
add_bullet(doc, "BREAKEVEN TIME EXIT: If SL has been moved to breakeven AND 15 minutes have passed with no progress toward TP, close at market. "
    "The slot is more valuable than a stalled trade.", bold_prefix="Rule 2")
add_bullet(doc, "SESSION CLOSE FORCED EXIT: Any scalping trade open at the end of its designated session window is force-closed at market. "
    "No scalper carries overnight. No scalper crosses into the next session. This is also halal compliance enforcement (no overnight positions). "
    "RVOL-LINKED TRIGGER: If session end is within 30 minutes AND RVOL is declining below 0.8x AND trade is in profit, take the profit immediately "
    "rather than waiting for hard session close.", bold_prefix="Rule 3")
add_bullet(doc, "MOMENTUM DEATH EXIT (RVOL-LINKED): If RVOL drops below 0.8x (market going dead) AND trade is open with no resolution, "
    "the momentum that created the setup is gone. Close at market. Rules 3 and 4 are linked — both use RVOL as the trigger signal.", bold_prefix="Rule 4")
add_note(doc,
    "ALL forced exits must be logged with a reason code in the Trading Journal. "
    "Reason codes: SESSION_END_CLOSE, TILT_PROFIT_TAKE, TILT_LOSS_ACCEPT, TIME_EXIT_BREAKEVEN, "
    "TIME_EXIT_LIMIT, RVOL_MOMENTUM_DEATH, NEWS_BLACKOUT_CLOSE. "
    "These codes feed into Workflow 3 (Performance Intelligence) for pattern analysis. "
    "High SESSION_END_CLOSE frequency on a strategy means its holding time assumption is wrong — "
    "flag for Workflow 2 (Improvement) in the next fortnightly cycle.",
    label="CODING AGENT NOTE")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-3  THE TILT — SESSION TRANSITION MECHANISM
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-3", "The Tilt — Universal Session Transition Mechanism")

add_sub_heading(doc, "S3-3.1  Definition")
add_body(doc,
    "The 'Tilt' is the system's term for a session transition moment — when one session is ending "
    "and another is beginning. It is a universal mechanism applied at every session boundary in the "
    "24-hour cycle. The Tilt is NOT a manual event. It is triggered automatically by the session "
    "template scheduler.")

add_sub_heading(doc, "S3-3.2  Tilt Trigger Conditions")
add_body(doc, "A Tilt is triggered when ANY of the following conditions are met:")
add_bullet(doc, "Session end time is T-5 minutes (5 minutes before the scheduled session close).")
add_bullet(doc, "RVOL for all active session instruments drops below 0.7x simultaneously.")
add_bullet(doc, "The Sentinel regime confidence drops below 0.40 for the outgoing session.")

add_sub_heading(doc, "S3-3.3  Tilt Sequence")
add_body(doc, "The Tilt executes the following sequence in order:")
add_bullet(doc, "LOCK: No new entries for any bot in the outgoing session pool.", level=0)
add_bullet(doc, "SIGNAL: Running trades receive the RVOL-linked profit signal. If in profit AND RVOL declining, close at market immediately. If at a loss, let existing SL/TP resolve it — do not force-close unless hard session end is reached.", level=0)
add_bullet(doc, "WAIT: 1-minute transition gap (system stabilisation, no actions).", level=0)
add_bullet(doc, "RE-RANK: DPR sub-agent scores the closing session bots. Queue is re-ordered for the incoming session.", level=0)
add_bullet(doc, "ACTIVATE: New session queue is unlocked. Bots cleared to trade in the incoming session.", level=0)
add_note(doc,
    "The Tilt is not just for London-NY transitions. It applies at: Sydney→Sydney-Tokyo, "
    "Sydney-Tokyo→Tokyo, Tokyo→London (with London Open Assault activation), "
    "London→Inter-session Cooldown, Cooldown→NY Open Assault, NY→Wind-down, Wind-down→Dead Zone, "
    "Dead Zone→Sydney. Every boundary has a Tilt.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-4  FULL 24-HOUR SESSION CYCLE ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-4", "Full 24-Hour Session Cycle Architecture")

add_sub_heading(doc, "S3-4.1  Session Template Design")
add_body(doc,
    "Rather than hard-coding separate logic per session, all sessions use a single configurable "
    "Session Template. The template has parameters that are set differently per session window. "
    "A coding agent implementing this should build ONE session template class with per-session "
    "configuration, not separate classes per session.")
add_body(doc, "Session Template parameters (vary per session):")
add_bullet(doc, "instruments_in_scope: list of currency pairs / instruments active in this session")
add_bullet(doc, "scalper_priority_weight: 0.0-1.0, proportion of queue slots allocated to scalper bots")
add_bullet(doc, "orb_priority_weight: 0.0-1.0, proportion of queue slots allocated to ORB bots (scalper + orb must sum to 1.0)")
add_bullet(doc, "kelly_fraction_modifier: multiplier applied to Kelly fraction during this session (1.0 = normal, 0.8 = conservative)")
add_bullet(doc, "max_concurrent_bots: maximum open positions at any time in this session (typically 5 for 2% risk each = 5% max exposure)")
add_bullet(doc, "house_money_threshold: daily PnL % required to trigger house money multiplier (lower = more aggressive in premium sessions)")
add_bullet(doc, "news_blackout_check: boolean — always true, but specified explicitly so coding agent does not omit it")
add_bullet(doc, "session_open_scalpers_active: boolean — true only for London Open and NY Open assault windows")
add_bullet(doc, "tilt_trigger_minutes: minutes before session end to initiate the Tilt sequence")

add_sub_heading(doc, "S3-4.2  Complete Daily Cycle Table")
make_table(doc,
    ["Window", "GMT Range", "Primary Type", "Key Instruments", "Scalper%", "ORB%", "Kelly Mod"],
    [
        ["Sydney Open",              "22:00-00:00", "ORB dominant",           "AUDUSD, AUDCAD, AUDNZD, AUDJPY", "30%", "70%", "0.80x"],
        ["Sydney-Tokyo Overlap",     "00:00-07:00", "Scalpers return",        "USDJPY, AUDJPY, GBPJPY",         "60%", "40%", "0.90x"],
        ["Tokyo Open",               "00:00-02:00", "ORB + Scalpers balanced","USDJPY, JPY crosses",            "50%", "50%", "0.90x"],
        ["Tokyo-London Overlap",     "07:00-09:00", "PREMIUM - Max aggression","EURUSD, GBPUSD, GBPJPY, EURJPY","80%", "20%", "1.00x"],
        ["London Open Assault",      "07:00-10:00", "Session Open Scalpers lead","EUR/GBP pairs",               "80%", "20%", "1.00x"],
        ["Inter-session Cooldown",   "10:00-13:00", "No trading - Eval only", "N/A",                           "N/A", "N/A", "N/A"],
        ["London-NY Overlap (PREMIUM)","13:00-16:00","Max aggression + House Money","All majors + Gold",        "70%", "30%", "1.00x"],
        ["NY Open Assault",          "13:00-15:30", "NY Session Scalpers lead","USD pairs, Gold",               "75%", "25%", "1.00x"],
        ["NY Wind-down",             "15:30-16:00", "Exit only, no new entries","N/A",                         "0%",  "0%",  "N/A"],
        ["Dead Zone",                "16:00-22:00", "Agentic only - no trading","N/A",                         "N/A", "N/A", "N/A"],
    ],
    col_widths=[1.8, 1.3, 1.8, 2.0, 0.7, 0.6, 0.9]
)
add_note(doc,
    "The Tokyo-London Overlap (07:00-09:00 GMT) is embedded within the London Open Assault window. "
    "When London opens, Tokyo is still running. JPY crosses (GBPJPY, EURJPY) are at peak volatility "
    "because both institutional populations are simultaneously active. Session Open Scalpers deployed "
    "during this overlap see the highest momentum of the day. This window is treated as PREMIUM "
    "even though it is not the London-NY overlap.")

add_sub_heading(doc, "S3-4.3  ORB Session Allocation Logic")
add_body(doc,
    "ORB (Opening Range Breakout) bots are session-masked. Each ORB variant is tagged with the "
    "session whose opening range it trades. ORB bots tagged for London Open trade the 07:00-07:30 "
    "range establishment, then watch for breakout. They are session-masked OFF by 11:00 GMT. "
    "ORB bots tagged for NY Open trade the 13:00-13:30 range, session-masked OFF by 17:00 GMT. "
    "ORB bots tagged for Tokyo Open trade the 00:00-00:30 range. "
    "In the London-NY premium overlap, ONLY NY ORB variants run (their range is establishing). "
    "London ORBs do NOT run in the afternoon. Scalpers dominate the premium window.")

add_sub_heading(doc, "S3-4.4  ORB in Asian Session — Special Allocation")
add_body(doc,
    "The Sydney Open (22:00-00:00 GMT) has low liquidity but clean directional movement. "
    "This is ORB-dominant territory because low volume = less noise = cleaner breakout signals. "
    "During Sydney Open, ORB bots are given 70% of the concurrent slot allocation. "
    "Capital allocation during Asian session remains at the same Kelly fraction as other sessions — "
    "house money threshold is NOT lowered for Asian sessions. The system does not increase aggression "
    "during low-liquidity windows.")

add_sub_heading(doc, "S3-4.5  News Kill Switch Integration (CRITICAL)")
add_body(doc,
    "Since the system now operates 24/7 across all sessions, the economic calendar news blackout "
    "must be integrated into EVERY session template. This was previously documented for London/NY "
    "sessions only. The rule now applies globally.")
add_bullet(doc, "Every session template checks the economic calendar before clearing any bot to trade.")
add_bullet(doc, "HIGH-IMPACT news events trigger a blackout window for the affected instrument pair: T-5 minutes before release to T+3 minutes after. No entries. Running trades are not force-closed but no new entries permitted.")
add_bullet(doc, "MEDIUM-IMPACT news: T-2 minutes to T+1 minute blackout.")
add_bullet(doc, "The SQS filter already partially handles this via spread spiking detection. The news calendar check is an additional explicit gate, not a replacement for SQS.")
add_bullet(doc, "News blackout applies even during Asian session. BOJ announcements, RBA decisions, RBNZ decisions all qualify as high-impact for their respective pair sessions.")
add_note(doc, "Coding agent: the economic calendar integration must be a shared service (similar to SVSS) that all session templates query. It should not be duplicated per session. One calendar service, queried by all.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-5  DEAD ZONE OPERATIONS & WORKFLOW 3
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-5", "Dead Zone Operations and Workflow 3 (Performance Intelligence)")

add_sub_heading(doc, "S3-5.1  Dead Zone Definition")
add_body(doc,
    "The Dead Zone is 16:00-22:00 GMT daily (6 hours). No new trading entries occur. "
    "All scalping positions must have been closed by the Tilt at 16:00 GMT. "
    "This window is the system's intelligence and analysis period. "
    "The Trading Journal is already being written to during live trading — the Dead Zone reads "
    "it, it does not write to it.")

add_sub_heading(doc, "S3-5.2  Workflow 3 — Performance Intelligence (Agent-Driven)")
add_body(doc,
    "DEFINITION: A Workflow in this system is any process that involves agents. If a process "
    "only involves system components executing automated logic, it is a system process, not a "
    "workflow. Workflow 3 is agent-driven because it requires intelligence synthesis, not just "
    "data aggregation.")
add_body(doc, "Workflow 3 runs during the Dead Zone every weekday. Steps:")
add_bullet(doc, "STEP 1 — EOD Report Generation: Trading Department agent reads the Trading Journal for the day. Produces a structured daily report: total trades, wins, losses, session-by-session breakdown, circuit breaker events, force-close events by reason code, net PnL, max drawdown, best performer, worst performer. Format mirrors a portfolio manager's end-of-day brief.", bold_prefix="16:15 GMT")
add_bullet(doc, "STEP 2 — Session-Based Performer Identification: Agent analyses performance by session window. Identifies bots that consistently outperform in specific sessions but underperform globally. Tags these bots as SESSION_SPECIALIST for their strong session. This tag elevates the bot to Tier 1 in its specialist session regardless of overall DPR composite score.", bold_prefix="16:45 GMT")
add_bullet(doc, "STEP 3 — DPR Score Update: System component (not agent) recalculates composite DPR scores for all bots based on the day's data. DPR score components: session win rate (25%), net PnL (30%), consistency score (equity curve smoothness, 20%), expected value per trade (25%). SESSION_SPECIALIST bots receive a tier override in their tagged session.", bold_prefix="17:00 GMT")
add_bullet(doc, "STEP 4 — Queue Re-ranking: System component re-orders the bot queue for tomorrow's sessions using updated DPR scores and SESSION_SPECIALIST tags.", bold_prefix="17:30 GMT")
add_bullet(doc, "STEP 5 — Fortnight Data Accumulation: Agent appends the daily report to the 14-day rolling intelligence file. When 14 days of data are accumulated, this file becomes the input to Workflow 4 (Weekend Update Cycle).", bold_prefix="18:00 GMT")

add_note(doc,
    "Bots are NOT updated during weekdays or during any active session. The Dead Zone collects "
    "and ranks — it does not modify strategy parameters or code. All bot updates happen on "
    "weekends only via Workflow 4. This is a hard rule to prevent unstable mid-week changes "
    "to a live system.")

add_sub_heading(doc, "S3-5.3  Session-Based Performer Tagging — Game Theory Element")
add_body(doc,
    "A bot may have an average overall DPR score but be the best-performing bot during a "
    "specific session. For example: a bot that underperforms in the afternoon NY session but "
    "consistently tops the London Open Assault leaderboard. This bot is tagged SESSION_SPECIALIST:LONDON_OPEN. "
    "During London Open Assault, it is always in Tier 1 regardless of composite DPR. "
    "This is a game theory concept: play your strengths, limit trade-offs.")
add_body(doc, "SESSION_SPECIALIST tag logic:")
add_bullet(doc, "Minimum 5 days of session-specific data required before tagging.")
add_bullet(doc, "Bot must rank in the top 20% for session win rate AND top 20% for session net PnL for that specific session.")
add_bullet(doc, "Tag persists until the bot underperforms in its specialist session for 3 consecutive days.")
add_bullet(doc, "A bot can hold multiple SESSION_SPECIALIST tags (e.g., LONDON_OPEN and NY_OPEN).")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-6  SURVIVORSHIP SELECTION LOOP (SSL) — CORRECTED ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-6", "Survivorship Selection Loop (SSL) — Corrected Architecture")

add_sub_heading(doc, "S3-6.1  Correct Language and Framing")
add_body(doc,
    "CRITICAL LANGUAGE NOTE FOR CODING AGENT: Do NOT use the phrase 'promoted to live' "
    "anywhere in the codebase for this mechanism. The correct language is 'cleared to trade' "
    "or 'given market access'. A bot does not change its operational state — it has always been "
    "built for live execution. It simply receives a queue window to trade. This distinction "
    "matters because 'promoted to live' implies a state change that would conflict with the "
    "AlphaForge paper-to-live gate (which is a separate, one-time, manually confirmed promotion).")

add_sub_heading(doc, "S3-6.2  SSL Flow")
add_body(doc, "The Survivorship Selection Loop is the daily rotation mechanism. It is automatic.")
add_code(doc, """
Bot loses 2 consecutive trades (personal book threshold)
        ↓
Circuit breaker fires → Bot enters Paper/Demo rotation
(Circuit breaker in paper trading = LOG ONLY, no quarantine effect)
        ↓
Live queue slot opens → Next ranked bot in DPR queue is CLEARED TO TRADE
(It was always built for live — it simply receives its queue window)
        ↓
Dedicated monitoring sub-agent (Trading Department) tracks paper rotation performance
Paper trading uses IDENTICAL system components as live trading — same Kelly, same SQS,
same Sentinel gate, same session masks. ONLY difference = virtual capital, no real orders.
        ↓
End of session: ALL bots scored by DPR (system component, not agent)
        ↓
Tomorrow's queue order = today's DPR composite scores + SESSION_SPECIALIST tag overrides
Best performer today → front of queue tomorrow
Quarantined bots in paper: their paper performance feeds DPR score but does NOT gate re-entry
        ↓
Weekend: Workflow 4 reviews fortnight data, decides parameter updates for consistently poor bots
""")

add_sub_heading(doc, "S3-6.3  Paper Trading Is Intelligence, Not a Gate")
add_body(doc,
    "A bot in paper trading is not being 're-evaluated for permission to return to live'. "
    "It was already fully validated before its first live trade (AlphaForge Workflow 1 → "
    "Workflow 2 → initial paper gate → manual one-time promotion). A quarantine event is "
    "market condition feedback, not a performance failure requiring re-certification. "
    "The paper trading data from a quarantined bot is valuable intelligence: if a bot "
    "continues to hit circuit breaker events even in paper trading (where the circuit "
    "breaker is logged not enforced), that is a strong negative signal for the DPR score "
    "and flags the bot for the fortnightly Workflow 4 improvement cycle.")

add_sub_heading(doc, "S3-6.4  Paper Trading Tiers")
add_body(doc, "The paper trading environment contains two populations:")
add_bullet(doc, "TIER 1 PAPER — Bots quarantined from live trading. These were previously validated and running live. They are temporarily rotating out. Their paper performance is held to a higher standard because they have a live history to compare against.", bold_prefix="Tier 1")
add_bullet(doc, "TIER 2 PAPER — Fresh bots from AlphaForge Workflow 2. They have never been live. These are the 'babies' fresh out of the coding and improvement pipeline. They must demonstrate at least 2 weeks of consistent paper performance before being eligible for the weekly live queue release.", bold_prefix="Tier 2")

add_sub_heading(doc, "S3-6.5  Consecutive Loss Threshold — Confirmed at 2")
add_labeled(doc, "Personal book (scalping)",  "2 consecutive losses → circuit breaker fires → paper rotation")
add_labeled(doc, "ORB bots",                  "3 consecutive losses (ORB trades less frequently, 2 would be too hair-trigger)")
add_labeled(doc, "Prop firm book (future)",   "Remains at 3 as configured in BotCircuitBreakerManager")
add_note(doc,
    "The rationale for 2 (not 5): at 52% WR, probability of 2 consecutive losses by chance = "
    "0.48² = 23%. Probability of 5 in a row = 0.48⁵ = 2.5%. Waiting for 5 consecutive losses "
    "means the bot is probably genuinely broken before it is pulled. At 2, the cost of a false "
    "positive (pulling a healthy bot) is one session off live — the queue fills the slot "
    "immediately. The benefit of catching a broken bot early justifies the threshold.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-7  DAILY PERFORMANCE RANKING (DPR) SYSTEM
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-7", "Daily Performance Ranking (DPR) System")

add_sub_heading(doc, "S3-7.1  Definition and Ownership")
add_labeled(doc, "What it is",      "An automated system component that scores all bots daily and determines queue order. NOT a workflow (no agents involved in the ranking calculation itself).")
add_labeled(doc, "Owner",           "Trading Department (system component). Portfolio Department monitors the output. FloorManager updates the Kanban board based on DPR output.")
add_labeled(doc, "Runs",            "Once per Dead Zone cycle (after Workflow 3 Step 3). Also recalculates at every Tilt event.")
add_labeled(doc, "Data source",     "Trading Journal (the canonical record of all trades, outcomes, timing, and reason codes).")

add_sub_heading(doc, "S3-7.2  DPR Composite Score Formula")
add_body(doc, "Each bot receives a composite score on a 0-100 scale:")
make_table(doc,
    ["Component", "Weight", "Description"],
    [
        ["Session Win Rate",        "25%", "Win rate for the most recent session window (not lifetime average). Rewards recent performance."],
        ["Net PnL",                 "30%", "Net PnL in dollars for the most recent session. Highest weight because absolute profit is the objective."],
        ["Consistency Score",       "20%", "Equity curve smoothness for the session. Calculated as 1 - (max_drawdown_in_session / total_gain_in_session). Higher = smoother curve = more consistent."],
        ["Expected Value / Trade",  "25%", "Net EV per trade for the session. EV = (win_rate × avg_win) - (loss_rate × avg_loss). Positive EV is required. Negative EV triggers a SESSION_CONCERN flag."],
    ],
    col_widths=[2.0, 0.8, 4.5]
)
add_note(doc, "SESSION_CONCERN flag: if a bot has negative EV for 3 consecutive sessions, it is flagged for priority review in the next Workflow 4 cycle (fortnightly). The flag does NOT quarantine the bot — it only elevates it in the review queue.")

add_sub_heading(doc, "S3-7.3  Queue Tier Remix Logic")
add_body(doc,
    "Queue order is NOT strict Tier 1 → Tier 2 → Tier 3. The confirmed remix approach is: "
    "Tier 1 bots interleaved with Tier 3 conditional re-entries, with Tier 2 solid performers "
    "filling the gaps. The reasoning: solid performers (Tier 2) will perform reliably regardless "
    "of queue position. Starting with the best (Tier 1), then testing a recovery candidate "
    "(Tier 3), then a solid performer (Tier 2), then another Tier 1, etc. — produces better "
    "learning data because you always have a control group (Tier 2) running alongside "
    "experimental slots (Tier 3).")
add_labeled(doc, "Example queue order", "T1, T3, T2, T1, T3, T2, T1, T2, T2, T2 (remaining slots)")
add_labeled(doc, "SESSION_SPECIALIST override", "Any bot with a SESSION_SPECIALIST tag for the current session is always placed in the first available Tier 1 slot for that session, regardless of composite DPR score.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-8  INTER-SESSION COOLDOWN — 3-HOUR INTELLIGENCE WINDOW
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-8", "Inter-Session Cooldown — 3-Hour Intelligence Window")

add_sub_heading(doc, "S3-8.1  Window Definition")
add_body(doc,
    "Between the London Open Assault close (~10:00 GMT) and the NY Open Assault start "
    "(13:00 GMT), there are approximately 3 hours. This is NOT idle time — it is the system's "
    "most important intra-day operational window. The Trading Department and Portfolio Department "
    "work together during this window. FloorManager updates the Kanban.")

add_sub_heading(doc, "S3-8.2  Cooldown Sequence")
add_bullet(doc, "10:00-10:30 GMT — Score all bots from London session. Produce DPR rankings. Flag quarantined bots.", bold_prefix="Step 1")
add_bullet(doc, "10:30-11:30 GMT — Monitoring sub-agent reviews paper trading data for quarantined bots. Review only — no parameter updates (weekday rule). Bots showing strong paper recovery are elevated to Tier 3 conditional re-entry for the NY session queue.", bold_prefix="Step 2")
add_bullet(doc, "11:30-12:40 GMT — Set NY Open queue order using DPR scores + SESSION_SPECIALIST tags + Tier remix logic. Tier 1 = best London performers. Tier 3 = first recovery candidates from quarantine review. AlphaForge does NOT release new variants to live during weekdays — new variants enter paper trading only.", bold_prefix="Step 3")
add_bullet(doc, "12:40-13:00 GMT — System health check. SQS pre-check for NY session instruments. Sentinel regime pre-confirmation for NY pairs. Session mask verification. 20-minute prep window before NY Open. The system walks into the NY Open Assault already warm and ranked.", bold_prefix="Step 4")
add_note(doc,
    "The hybrid queue for NY Open: the bot that topped the London Open leaderboard goes first. "
    "The best-performing paper recovery candidate (highest paper DPR score) goes second. "
    "Then standard Tier 1/3/2 remix. This produces the highest-quality roster for the premium "
    "London-NY overlap window that begins at 13:00.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-9  CONFIRMED WORKFLOW INVENTORY
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-9", "Confirmed Workflow Inventory (All Workflows)")

add_body(doc,
    "DEFINITION LOCK: In this system, a Workflow is any process that involves AI agents. "
    "If a process only involves system components executing automated logic without agents, "
    "it is a SYSTEM PROCESS, not a Workflow. DPR scoring is a system process. "
    "The Tilt is a system process. The following are the confirmed Workflows.")

add_sub_heading(doc, "S3-9.1  Workflow 1 — AlphaForge (Strategy Creation)")
add_body(doc, "Purpose: Turn raw strategy ideas (from any source) into a working, backtested EA.")
add_code(doc, """
INPUT: Strategy seed (video ingest, Firecrawl scrape, article, indicator combination, manual design)
  ↓
Research Department agent: extracts trading logic, creates SDD (Strategy Design Document)
  ↓
Coding: EA implementation in MQL5
  ↓
Compiling: MT5 compile check, syntax validation
  ↓
Backtesting: historical backtest across multiple market conditions
  ↓
Monte Carlo simulation: 1000+ permutations to test robustness
  ↓
Report: performance metrics, win rate, drawdown, R:R, edge score
  ↓
OUTPUT: EA file + performance report

WORKFLOW 1 ENDS HERE. Refinement is handled by Workflow 2, not here.
""")

add_sub_heading(doc, "S3-9.2  Workflow 2 — Improvement Loop (Refinement and Paper Gate)")
add_body(doc, "Purpose: Continuously improve strategy parameters and validate via paper trading.")
add_code(doc, """
INPUT: EA from Workflow 1 (new bot) OR EA flagged for improvement by Workflow 3/4 (existing bot)
  ↓
Refinement: parameter optimisation, indicator tuning
  ↓
Walk-Forward Analysis (WFA): out-of-sample validation
  ↓
Distillation: extract best-performing parameter set from WFA
  ↓
BRANCH A — NEW BOT: Enter paper trading gate (minimum 2 weeks, Tier 2 paper population)
           If paper performance passes → eligible for weekly live queue release
           If paper performance fails → loop back to Refinement
  ↓
BRANCH B — EXISTING BOT IMPROVEMENT: Create new VARIANT. Variant enters Tier 2 paper trading.
           Old bot continues in live (degraded DPR ranking, fewer queue slots).
           After 2 weeks: if new variant outperforms old bot in paper → old bot exits live gracefully,
           new variant enters at Tier 3 conditional. If variant underperforms → discard variant,
           investigate further, try another cycle.
  ↓
OUTPUT: Validated EA variant ready for live queue OR looped back for further improvement
""")

add_sub_heading(doc, "S3-9.3  Workflow 3 — Performance Intelligence (Dead Zone, Agent-Driven)")
add_body(doc, "Purpose: Daily intelligence synthesis and reporting. Runs every weekday in the Dead Zone.")
add_code(doc, """
INPUT: Trading Journal (read-only during weekdays)
  ↓
EOD Report agent (Trading Department): structured daily performance brief
  ↓
Session-based performer identification: SESSION_SPECIALIST tag assignments
  ↓
DPR score update (SYSTEM PROCESS, not agent — automated calculation)
  ↓
Fortnight data accumulation: appends daily report to 14-day rolling intelligence file
  ↓
OUTPUT: Daily EOD report + updated DPR scores + SESSION_SPECIALIST tags + fortnight file

NOTE: No bot parameters are updated. No code changes. Read and rank only.
""")

add_sub_heading(doc, "S3-9.4  Workflow 4 — Weekend Update Cycle (Agent-Driven)")
add_body(doc, "Purpose: Weekly improvement cycle based on fortnightly intelligence data.")
add_code(doc, """
INPUT: 14-day rolling intelligence file from Workflow 3
  ↓
Planning session (Trading Dept + Portfolio Dept agents): identify bots flagged SESSION_CONCERN,
underperformers, and improvement opportunities. Produce weekend action plan.
  ↓
Friday night: analysis and planning (markets closed)
  ↓
Saturday: Workflow 2 triggered for flagged bots — refinement, WFA, new variants
  ↓
Saturday-Sunday: testing, compilation, backtesting of new variants
  ↓
Sunday night: pre-market preparation, new paper trading slots opened for new variants,
Monday queue ranked and ready.
  ↓
OUTPUT: Updated bot variants in paper trading, clean Monday queue
""")

add_sub_heading(doc, "S3-9.5  Strategy Intelligence Harvester (SIH) — Input Pipeline")
add_body(doc,
    "The SIH is the data input layer that feeds Workflow 1. It is not a standalone workflow "
    "— it is the source of seeds that enter Workflow 1.")
add_bullet(doc, "VIDEO INGEST: Already implemented as a tool in the system. Ingests YouTube videos containing trading strategy content. Produces unbiased extracted information. Research Department agent then processes this into an SDD for Workflow 1.")
add_bullet(doc, "FIRECRAWL SCRIPT (MT5 COMMUNITY): A script (not an agent) scrapes the MT5 freelance community section for job postings. People paying for EA development reveal which strategies have demand. Research Department agent filters for scalping-logic strategies. Output feeds Workflow 1.")
add_bullet(doc, "INDICATOR COMBINATIONS: Agents (during Dead Zone or weekend) can propose new indicator combinations as strategy seeds. These can be generated from articles, trading textbooks, or manual design sessions. They enter Workflow 1 as seeds.")
add_bullet(doc, "FIRECRAWL SCRIPT (ARTICLES): An existing Firecrawl script for article scraping. Can be extended to target MT5 community with minimal changes.")
add_note(doc, "DPR is NOT a step in the SIH pipeline. DPR is a system process that runs independently. The SIH → Workflow 1 → Workflow 2 → Paper Gate → Live Queue pipeline is linear with branches. DPR runs continuously in parallel, scoring whatever is in the live pool.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-10  SESSION OPEN ASSAULT ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-10", "Session Open Assault Architecture")

add_sub_heading(doc, "S3-10.1  Concept")
add_body(doc,
    "The Session Open Assault is a deliberate concentration of session-open-specialist scalper "
    "bots during the first 60-180 minutes of high-liquidity session openings. The objective is "
    "to capture the maximum directional energy of the opening push before the market settles "
    "into range or drift. Session Open Scalpers are a dedicated bot family in the scalping "
    "strategy library — they are distinct from general momentum scalpers.")

add_sub_heading(doc, "S3-10.2  London Open Assault (07:00-10:00 GMT)")
add_labeled(doc, "Target pairs",         "EURUSD, GBPUSD, EURGBP, GBPJPY — highest pip movement in first 90 minutes")
add_labeled(doc, "Minute 0-90",          "Maximum aggression. Session Open Scalpers fully deployed. Momentum scalpers alongside them. Sentinel confirms regime direction at 06:55 GMT (5 minutes pre-open).")
add_labeled(doc, "Minute 90-180",        "Wind down momentum scalpers. ORB logic takes over as the market begins to define its range.")
add_labeled(doc, "ORB during assault",   "Lower queue priority than scalpers but NOT excluded. If an ORB setup fires and a slot is free, it trades. If all 5 concurrent slots are scalpers, ORB waits. Priority queue: scalpers = 1-5, ORBs = 6-8.")

add_sub_heading(doc, "S3-10.3  NY Open Assault (13:00-16:00 GMT — Premium Window)")
add_labeled(doc, "Target pairs",         "EURUSD, GBPUSD, USDCAD, XAUUSD (Gold)")
add_labeled(doc, "House money",          "London-NY overlap triggers house money multiplier at +4% daily PnL threshold (lower than the normal +8%). During this window the system deploys more aggressively because liquidity is highest.")
add_labeled(doc, "NY Opening Range ORBs","ORB bots specifically designed for the NY 13:00-13:30 range establishment are valid in this window. London ORBs are NOT active past 11:00 GMT.")
add_labeled(doc, "Minute 0-90 (13:00-14:30)", "Maximum NY Open Assault. Session Open Scalpers + NY ORBs active.")
add_labeled(doc, "Minute 90-150 (14:30-16:00)", "Wind down. Premium overlap closes at 16:00 with the Tilt.")

add_sub_heading(doc, "S3-10.4  The 3-Hour Inter-Session Intelligence Window")
add_body(doc,
    "Between London Open close (~10:00) and NY Open start (13:00) there are approximately 3 hours. "
    "This gap is used for the Inter-Session Cooldown (see Section S3-8). The NY Open queue is "
    "set using London session performance data, creating a meritocratic performance-tiered roster "
    "for the premium window. Best London performer → first in NY queue.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-11  HOUSE OF MONEY — CORRECTED THRESHOLDS
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-11", "House of Money — Corrected Threshold Architecture")

add_sub_heading(doc, "S3-11.1  Problem with Current Implementation")
add_body(doc,
    "The current EnhancedGovernor has the preservation mode trigger at -3% daily PnL. "
    "On a $200 account doing 30+ trades at 2% risk, -3% = $6 = approximately 1.5 losing trades. "
    "This is normal variance for a scalping system, not a genuine problem day. The -3% threshold "
    "causes the system to constantly throttle itself on variance, suppressing compounding.")

add_sub_heading(doc, "S3-11.2  Confirmed Corrected Thresholds")
make_table(doc,
    ["Parameter", "Current Value", "Corrected Value", "Rationale"],
    [
        ["House money trigger",         "+5% daily PnL",  "+8% daily PnL",         "Let the edge play out before boosting. At 2% risk/30 trades, +8% = ~4 net wins above average."],
        ["House money multiplier up",   "1.5x",           "Scales 1.25x → 2.0x",   "Continuous scaling (not binary). Formula: multiplier = 1.0 + (pnl_pct / 0.10), capped at 2.5x."],
        ["Preservation trigger",        "-3% daily PnL",  "-10% daily PnL",         "-10% = 5 net losses at 2% risk. This is genuine system malfunction, not variance."],
        ["Preservation multiplier",     "0.5x",           "0.5x (unchanged)",       "The multiplier is correct. Only the trigger threshold was wrong."],
        ["London-NY overlap trigger",   "N/A",            "+4% daily PnL",          "Lower threshold in premium session. Deploy faster when liquidity is highest."],
        ["London-NY overlap multiplier","N/A",            "1.75x → 2.5x scaling",   "More aggression in the premium window but with the overlap consecutive loss protection."],
        ["Overlap protection",          "N/A",            "3 consecutive losses in overlap → reset multiplier to 1.0x for that session only", "Prevents a bad overlap run from spiralling."],
        ["Session reset",               "Daily",          "Daily at 00:00 GMT (start of Sydney cycle)", "House money state resets at the start of each full trading day cycle."],
    ],
    col_widths=[2.0, 1.5, 1.8, 2.4]
)

add_sub_heading(doc, "S3-11.3  Continuous Scaling Formula")
add_body(doc, "Replace the binary 1.0x / 1.5x flip with continuous scaling:")
add_code(doc, """
def _update_house_money_effect(self, current_balance: float) -> None:
    pnl_pct = (current_balance - self._daily_start_balance) / self._daily_start_balance

    if pnl_pct > 0.08:  # Above +8%: house money active, continuous scaling
        multiplier = 1.0 + (pnl_pct / 0.10)
        multiplier = min(multiplier, 2.5)   # Hard ceiling at 2.5x
    elif pnl_pct > 0.04:  # +4% to +8%: early house money
        multiplier = 1.25
    elif pnl_pct < -0.10:  # Below -10%: preservation mode
        multiplier = 0.5
    else:
        multiplier = 1.0  # Normal range

    # London-NY overlap override (applied on top if in premium session)
    if self._in_london_ny_overlap():
        if pnl_pct > 0.04:  # Lower trigger in premium session
            multiplier = max(multiplier, 1.0 + (pnl_pct / 0.08))
            multiplier = min(multiplier, 2.5)
        if self._overlap_consecutive_losses >= 3:
            multiplier = 1.0  # Reset on 3 consecutive overlap losses

    self.house_money_multiplier = multiplier
""")
add_note(doc, "The _in_london_ny_overlap() method checks if current UTC time is between 13:00 and 16:00 GMT. _overlap_consecutive_losses is a session-scoped counter reset at 16:00 GMT daily.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-12  ENHANCED KELLY — ANALYSIS AT CONFIRMED R:R RATIOS
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-12", "Enhanced Kelly — Analysis at Confirmed R:R Ratios")

add_sub_heading(doc, "S3-12.1  Kelly Fraction at 1:2 R:R")
add_code(doc, """
win_rate = 0.52, payoff_ratio = 2.0

# Base Kelly
f* = ((2 + 1) × 0.52 - 1) / 2 = (1.56 - 1) / 2 = 0.56 / 2 = 0.28 = 28%

# After Layer 1 (Half-Kelly, kelly_fraction=0.50 in EnhancedKellyConfig):
f_layer1 = 0.28 × 0.50 = 0.14 = 14%

# After Layer 2 (Hard cap, max_risk_pct=0.02):
f_final = min(0.14, 0.02) = 0.02 = 2%

# Result: System always caps at 2% per trade at 1:2 R:R, 52% WR
# On $200 account: risk per trade = $200 × 0.02 = $4
""")

add_sub_heading(doc, "S3-12.2  Kelly Fraction at 1:3 R:R")
add_code(doc, """
win_rate = 0.52, payoff_ratio = 3.0

# Base Kelly
f* = ((3 + 1) × 0.52 - 1) / 3 = (2.08 - 1) / 3 = 1.08 / 3 = 0.36 = 36%

# After Layer 1 (Half-Kelly):
f_layer1 = 0.36 × 0.50 = 0.18 = 18%

# After Layer 2 (Hard cap):
f_final = min(0.18, 0.02) = 0.02 = 2%

# Both R:R targets hit the 2% hard cap. The cap is the binding constraint.
# This is correct and intentional — 2% is conservative by design at startup.
# As equity grows and confidence builds, max_risk_pct can be reviewed.
""")

add_sub_heading(doc, "S3-12.3  Early Phase Fallback")
add_body(doc,
    "The EnhancedKellyConfig.fallback_risk_pct = 0.01 (1%) is used when a bot has fewer than "
    "min_trade_history = 30 validated trades. In the early phase (first few weeks of live trading), "
    "all bots will be in fallback mode: 1% risk = $2 per trade on a $200 account. "
    "Once 30 trades are accumulated, the system switches to the Kelly-calculated fraction "
    "(which immediately hits the 2% cap at these R:R ratios). "
    "There is no manual switch required — this is automatic in EnhancedKellyCalculator.")

add_sub_heading(doc, "S3-12.4  Kelly Is Already Optimised for 1:2 R:R")
add_body(doc,
    "The EnhancedKellyCalculator already performs fee-adjusted Kelly calculation. "
    "avg_win and avg_loss are adjusted for commission and spread before the Kelly formula runs. "
    "This means the 2% output already accounts for trading costs — do NOT deduct fees "
    "separately from EV calculations. The Kelly output IS the fee-adjusted risk fraction.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-13  SEVEN SCALPING LOGIC FAMILIES
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-13", "Seven Scalping Logic Families")

add_body(doc,
    "IMPORTANT CLARIFICATION: 'Scalping' describes a holding period (M1/M5 timeframes, "
    "typically under 30 minutes). It does NOT describe a specific entry logic. The following "
    "seven families each use different entry logic but all share the scalping holding period "
    "constraint. Each family can have hundreds of variants (different instruments, parameter "
    "sets, indicator combinations). All variants share the same SL/TP hybrid mechanics (S3-2), "
    "the same Kelly sizing (S3-12), and the same circuit breaker (2 consecutive losses).")

make_table(doc,
    ["Family", "Entry Logic", "Pool", "Best Session", "Notes"],
    [
        ["Momentum Scalpers",       "Enter in direction of short-term momentum (dual EMA crossover, RSI momentum break, VWAP momentum deviation)",
         "Long / Short", "London & NY Open Assault", "Highest frequency. Most trades per session."],
        ["Mean Reversion Scalpers", "Fade moves that overshoot the mean (Bollinger Band bounce, RSI extremes, VWAP deviation fade)",
         "Neutral", "All sessions", "Neutral Pool — do not care about trend direction."],
        ["Micro Breakout Scalpers", "Enter on break of very short-term M1/M5 consolidation range. NOT session ORB — this is a 15-30 min micro range.",
         "Long / Short", "London & NY Open", "Distinct from ORB. Shorter timeframe, smaller range."],
        ["Liquidity Grab Scalpers", "Enter after a stop-hunt spike — price wicks below key M5 low, sweeps stops, then reverses. Enter the reversal.",
         "Neutral / Directional", "London & NY Overlap", "High-probability if timed well. Requires volume confirmation."],
        ["Session Open Scalpers",   "Specifically designed for the first 15-90 minutes of London or NY open. Exploit the initial directional push.",
         "Long / Short", "London 07:00-09:00 / NY 13:00-14:30", "Session-masked — ONLY active in their designated open window."],
        ["Volume Confirmation Scalpers", "Enter only when RVOL > 1.2x (confirmed real participation). Wrapper logic around any of the above families.",
         "Long / Short / Neutral", "All high-liquidity windows", "RVOL gate is the distinguishing feature. Uses SVSS data."],
        ["Spread Compression Scalpers", "Enter during the brief window after a spread spike normalises back to baseline. SQS detects this moment.",
         "Long / Short", "Session opens (spread spikes at open then compresses)", "SQS filter is the entry trigger. Exploits the liquidity surge after spread compression."],
    ],
    col_widths=[1.7, 2.5, 1.0, 1.5, 2.0]
)

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-14  VARIANT REPLACEMENT CYCLE
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-14", "Variant Replacement Cycle — Graceful Deprecation")

add_body(doc,
    "When a live bot is consistently underperforming and flagged by Workflow 3 for improvement, "
    "Workflow 2 creates a new VARIANT. The old bot is NOT abruptly removed. This is graceful "
    "deprecation — the system maintains live coverage while the replacement is being validated.")
add_code(doc, """
Live bot B_v1 is consistently underperforming (SESSION_CONCERN flag for 3 consecutive sessions)
        ↓
Workflow 4 (weekend) triggers Workflow 2 for B_v1
        ↓
Workflow 2 creates B_v2 (improved variant) → B_v2 enters Tier 2 paper trading
        ↓
B_v1 continues in live trading but its DPR score is reduced → it gets fewer queue slots
        ↓
B_v2 paper trades for minimum 2 weeks (Tier 2 paper population)
        ↓
COMPARISON: B_v2 paper performance vs B_v1 live performance
        ↓
IF B_v2 outperforms B_v1:
    → B_v1 is gracefully retired (DPR score drops to 0, naturally exits queue)
    → B_v2 enters live at Tier 3 conditional
    → Bot tagging system records B_v1 → B_v2 genealogy
IF B_v2 underperforms B_v1:
    → B_v2 is discarded or looped back to Workflow 2 for further refinement
    → B_v1 continues (it is still the best available version)
""")
add_note(doc,
    "The bot tagging system records the genealogy of every variant. B_v1 → B_v2 → B_v3 etc. "
    "This allows the agentic system to track improvement trajectories and identify strategy "
    "families that consistently improve vs those that plateau. Plateau strategies are candidates "
    "for retirement and replacement with entirely new AlphaForge Workflow 1 strategies.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-15  WEEKEND CYCLE AND BOT UPDATE RULES
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-15", "Weekend Cycle and Bot Update Rules")

add_sub_heading(doc, "S3-15.1  The No-Update-During-Week Rule")
add_body(doc,
    "HARD RULE: Bot parameters are NOT updated during weekdays or during any active session. "
    "The Dead Zone (16:00-22:00 GMT) collects intelligence and ranks — it does not modify "
    "strategy parameters or code. This rule exists to prevent unstable mid-week changes to "
    "a live system. A single bad parameter change on a Tuesday could corrupt the rest of the "
    "week's data.")
add_body(doc, "During weekdays:")
add_bullet(doc, "AlphaForge can PRODUCE new variants and send them to paper trading at any time (this does not affect the live pool).")
add_bullet(doc, "Workflow 3 reads and ranks (no writes to strategy parameters).")
add_bullet(doc, "The monitoring sub-agent reviews quarantined bots (review only, no updates).")
add_body(doc, "During weekends:")
add_bullet(doc, "Workflow 4 executes. Workflow 2 runs for flagged bots.")
add_bullet(doc, "New variants are tested, refined, and prepared for Monday paper trading entry.")
add_bullet(doc, "The live queue for Monday is ranked and pre-loaded.")

add_sub_heading(doc, "S3-15.2  Fortnightly Data as Update Basis")
add_body(doc,
    "Daily data is insufficient for meaningful parameter updates. The minimum data window "
    "for bot parameter changes is 14 days (fortnightly). Workflow 4 uses the 14-day rolling "
    "intelligence file from Workflow 3 as its primary input. Updates based on fewer than "
    "14 days of data risk over-fitting to short-term noise.")

add_sub_heading(doc, "S3-15.3  New Variants Released on Monday Only")
add_body(doc,
    "New AlphaForge-produced variants that have completed the 2-week paper trading gate "
    "are released into the live queue on Monday morning only — not mid-week. This creates "
    "a clean weekly cadence: evaluate on Friday, update over the weekend, fresh roster "
    "deploys Monday.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-16  SYSTEM COMPLEXITY ASSESSMENT
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-16", "System Complexity Assessment")

add_sub_heading(doc, "S3-16.1  Complexity Score: 9.1 / 10")
add_body(doc,
    "For context: a retail trader implementing a single bot on MT5 is complexity 3. "
    "A quant fund running a multi-strategy portfolio with regime detection is complexity 8. "
    "QUANTMINDX sits above institutional mid-tier because of three specific components that "
    "do not exist in combination elsewhere: the agentic self-improvement loop (AlphaForge + "
    "Workflow 2 running agents 24/7), the Survivorship Selection Loop with DPR-based daily "
    "rotation, and the Correlation Sensor using Marchenko-Pastur Random Matrix Theory (PhD-level "
    "quantitative finance).")

add_sub_heading(doc, "S3-16.2  Realistic Timeline to June 2026 Live Launch")
add_body(doc, "June 2026 is the confirmed target for live trading. The timeline is achievable because:")
add_bullet(doc, "AlphaForge agents work 24/7. Strategy generation is not the bottleneck.")
add_bullet(doc, "The bottleneck is the 2-week paper trading gate. With 10-15 strategies in parallel paper trading, you validate 5-8 per week.")
add_bullet(doc, "Target: 45 validated bots in the live queue before first live trade.")
add_bullet(doc, "At 6 validations per week from early April: 45 ÷ 6 = 7.5 weeks = mid-May paper complete.")
add_bullet(doc, "Add 3 weeks for system hardening, bug discovery from paper trading, pre-live checklist: early to mid June.")
add_bullet(doc, "The video ingest + MT5 Firecrawl pipeline accelerates ideation to near-zero time.")

add_sub_heading(doc, "S3-16.3  Can the System Perform as Designed?")
add_body(doc,
    "The math is sound. The edge sources are real. The architecture is coherent. The compound "
    "growth numbers are theoretical ceilings that the system approaches asymptotically as the "
    "bot library grows and strategies validate. The complexity is not a risk to the plan — "
    "the complexity IS the plan. Every layer (regime detection, Kelly, house money, correlation "
    "sensor, circuit breakers, SSL, DPR, session templates) is a protection layer that makes "
    "the compound growth curve steeper and more durable. A simpler system would start faster "
    "and blow up sooner.")
add_note(doc,
    "FOR CODING AGENT: The system is large but each component has a defined boundary. "
    "Build each component to its interface specification. Do not let components bleed into "
    "each other's responsibilities. The FloorManager orchestrates. Departments own their domains. "
    "Workflows involve agents. System processes involve code. Kanban tracks state. "
    "Trading Journal is the source of truth for performance data. Everything else reads from it.")

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# S3-17  OPEN ITEMS AND NEXT PLANNING SESSION AGENDA
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "S3-17", "Open Items and Next Planning Session Agenda")

add_sub_heading(doc, "Items to Confirm / Design in Next Session")
add_bullet(doc, "Circuit breaker daily trade limit per bot family (currently global constant = 20, needs per-family configuration wired into check_allowed() method in BotCircuitBreakerManager).")
add_bullet(doc, "Session template class design — confirm the parameter schema and implement the shared session template object.")
add_bullet(doc, "Economic calendar service — design as shared service queried by all session templates (similar to SVSS architecture).")
add_bullet(doc, "Cloudzy VPS sizing — confirm RAM and CPU are sufficient for peak concurrent load at Tokyo-London overlap (highest concurrent session overlap).")
add_bullet(doc, "ORB bot session masking — implement explicit session tags on ORB variants so London ORBs cannot fire after 11:00 GMT and NY ORBs activate at 13:00 GMT only.")
add_bullet(doc, "SESSION_SPECIALIST tag persistence and storage — define in data model.")
add_bullet(doc, "Workflow 3 and 4 agent team definitions — which Department sub-agents own each step.")
add_bullet(doc, "Bot tagging system integration with variant replacement cycle — define genealogy data structure.")

# ── FINAL PAGE: session close marker ────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
shade_paragraph(p, "375A32")
r = p.add_run("  SESSION 3 ADDENDUM COMPLETE — March 21, 2026")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = WHITE
r.font.name = "Calibri"

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(ADDENDUM_PATH)
print(f"Saved: {ADDENDUM_PATH}")
